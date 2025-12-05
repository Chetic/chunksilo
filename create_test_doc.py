#!/usr/bin/env python3
"""Create test DOCX, Markdown, and TXT documents for testing."""
from docx import Document
from pathlib import Path

# Create data directory
data_dir = Path('data')
data_dir.mkdir(exist_ok=True)

# ============================================================================
# Create DOCX document
# ============================================================================
doc = Document()

# Add content
doc.add_heading('Sample Documentation', 0)
doc.add_paragraph('This is a sample DOCX document for testing the retrieval-augmented documentation system.')

doc.add_heading('Introduction', 1)
doc.add_paragraph(
    'This document describes a sample retrieval-augmented system implementation. '
    'The system is designed to index and query documentation using semantic search over vector embeddings.'
)

doc.add_heading('Key Features', 1)
doc.add_paragraph('1. Document Indexing: The system can index PDF, DOCX, Markdown, and TXT documents from a local directory.')
doc.add_paragraph('2. Semantic Search: Uses vector embeddings to find relevant document chunks based on semantic similarity.')
doc.add_paragraph('3. MCP Integration: Exposes retrieval capabilities via the Model Context Protocol (MCP).')
doc.add_paragraph('4. Answer Synthesis by the Client: A separate LLM client (e.g., Continue) uses the retrieved chunks to generate answers.')

doc.add_heading('Architecture', 1)
doc.add_paragraph('The system consists of three main components:')
doc.add_paragraph('- Ingestion Pipeline: Processes documents and builds a vector index using an embedding model')
doc.add_paragraph('- MCP Server: Handles queries and returns relevant document chunks with metadata')
doc.add_paragraph('- LLM Client (e.g., Continue): Calls the MCP tool to retrieve chunks and synthesizes answers from them')

doc.add_heading('Usage', 1)
doc.add_paragraph('To use the system:')
doc.add_paragraph('1. Place PDF, DOCX, Markdown, or TXT files in the data directory')
doc.add_paragraph('2. Run the ingestion script to build the index')
doc.add_paragraph('3. Start the MCP server')
doc.add_paragraph('4. Use an MCP-aware client (such as Continue) to call the retrieval tool and generate answers from the returned chunks')

# Save DOCX
doc.save(data_dir / 'sample_document.docx')
print(f"Created {data_dir / 'sample_document.docx'}")

# ============================================================================
# Create Markdown document
# ============================================================================
markdown_content = """# Sample Documentation

This is a sample Markdown document for testing the retrieval-augmented documentation system.

## Introduction

This document describes a sample retrieval-augmented system implementation. The system is designed to index and query documentation using semantic search over vector embeddings.

## Key Features

1. **Document Indexing**: The system can index PDF, DOCX, Markdown, and TXT documents from a local directory.
2. **Semantic Search**: Uses vector embeddings to find relevant document chunks based on semantic similarity.
3. **MCP Integration**: Exposes retrieval capabilities via the Model Context Protocol (MCP).
4. **Answer Synthesis by the Client**: A separate LLM client (e.g., Continue) uses the retrieved chunks to generate answers.

## Architecture

The system consists of three main components:

- **Ingestion Pipeline**: Processes documents and builds a vector index using an embedding model
- **MCP Server**: Handles queries and returns relevant document chunks with metadata
- **LLM Client** (e.g., Continue): Calls the MCP tool to retrieve chunks and synthesizes answers from them

## Usage

To use the system:

1. Place PDF, DOCX, Markdown, or TXT files in the data directory
2. Run the ingestion script to build the index
3. Start the MCP server
4. Use an MCP-aware client (such as Continue) to call the retrieval tool and generate answers from the returned chunks

## Markdown Support

This document demonstrates that the system now supports Markdown files. Markdown is a lightweight markup language that is widely used for documentation, README files, and technical writing. The system can process standard Markdown syntax including:

- Headers (like this section)
- **Bold** and *italic* text
- Lists (both ordered and unordered)
- Code blocks
- Links and references
"""

# Save Markdown
markdown_path = data_dir / 'sample_document.md'
markdown_path.write_text(markdown_content)
print(f"Created {markdown_path}")

