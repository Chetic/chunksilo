#!/usr/bin/env python3
# SPDX-License-Identifier: Apache-2.0
"""
Indexing pipeline for building a RAG index from PDF, DOCX, DOC, Markdown, and TXT documents.
Supports incremental indexing using a local SQLite database to track file states.
"""
import argparse
import hashlib
import itertools
import json
import logging
import os
import signal
import sqlite3
import sys
import threading
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Optional, Iterator, Set, Any, Tuple

from docx import Document

from llama_index.core import (
    VectorStoreIndex,
    StorageContext,
    Settings,
    SimpleDirectoryReader,
    Document as LlamaIndexDocument,
    load_index_from_storage,
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import MetadataMode
from llama_index.embeddings.fastembed import FastEmbedEmbedding

# Load configuration from config.yaml
from . import cfgload
from .cfgload import load_config
_config = load_config()

# Configuration from config.yaml
STORAGE_DIR = Path(_config["storage"]["storage_dir"])
STATE_DB_PATH = STORAGE_DIR / "ingestion_state.db"

# Stage 1 (embedding/vector search) configuration
RETRIEVAL_EMBED_MODEL_NAME = _config["retrieval"]["embed_model_name"]

# Stage 2 (FlashRank reranking, CPU-only, ONNX-based) configuration
RETRIEVAL_RERANK_MODEL_NAME = _config["retrieval"]["rerank_model_name"]

# Shared cache directory for embedding and reranking models
RETRIEVAL_MODEL_CACHE_DIR = Path(_config["storage"]["model_cache_dir"])

# BM25 index directory for file name matching
BM25_INDEX_DIR = STORAGE_DIR / "bm25_index"

# Heading store for document headings (stored separately to avoid metadata size issues)
HEADING_STORE_PATH = STORAGE_DIR / "heading_store.json"

# Metadata exclusion configuration
# These keys are excluded from the embedding text to save tokens and avoid length errors
EXCLUDED_EMBED_METADATA_KEYS = [
    "line_offsets",      # Large integer array, primary cause of length errors
    "document_headings", # Heading hierarchy array with positions, excluded like line_offsets
    "heading_path",      # Pre-computed heading hierarchy, stored separately to save chunk space
    "file_path",         # redundant with file_name/source, strict path less useful for semantic similarity
    "source",            # often same as file_path
    "creation_date",     # temporal, not semantic
    "last_modified_date",# temporal, not semantic
    "doc_ids",           # internal tracking
    "hash",              # internal tracking
]

# These keys are excluded from the LLM context to save context window
EXCLUDED_LLM_METADATA_KEYS = [
    "line_offsets",      # LLM needs text content, not integer map
    "hash",              # internal tracking
    "doc_ids",           # internal tracking
    "file_path",         # usually redundant if file_name is present
    "source",            # usually redundant
]

logger = logging.getLogger(__name__)


# Default file type patterns
DEFAULT_INCLUDE_PATTERNS = ["**/*.pdf", "**/*.md", "**/*.txt", "**/*.docx", "**/*.doc"]

# Default directory/file exclusion patterns — skip common noise directories
DEFAULT_EXCLUDE_PATTERNS = [
    "**/.git/**",
    "**/node_modules/**",
    "**/__pycache__/**",
    "**/.venv/**",
    "**/venv/**",
    "**/.tox/**",
    "**/.mypy_cache/**",
    "**/.pytest_cache/**",
    "**/.eggs/**",
    "**/*.egg-info/**",
    "**/.DS_Store",
]


@dataclass
class DirectoryConfig:
    """Configuration for a single source directory."""
    path: Path
    enabled: bool = True
    include: List[str] = field(default_factory=lambda: DEFAULT_INCLUDE_PATTERNS.copy())
    exclude: List[str] = field(default_factory=lambda: DEFAULT_EXCLUDE_PATTERNS.copy())
    recursive: bool = True
    case_sensitive: bool = False


@dataclass
class IndexConfig:
    """Complete indexing configuration."""
    directories: List[DirectoryConfig]
    chunk_size: int = 512
    chunk_overlap: int = 50


def load_index_config() -> IndexConfig:
    """Load indexing configuration from config.yaml.

    Raises:
        ValueError: If config is invalid
    """
    indexing_config = _config.get("indexing", {})

    if not indexing_config.get("directories"):
        raise ValueError(
            "Config must have at least one directory in 'indexing.directories'.\n"
            "Please update config.yaml with your directory configuration.\n"
            "Example:\n"
            "indexing:\n"
            "  directories:\n"
            '    - "./data"\n'
            "  chunk_size: 512\n"
            "  chunk_overlap: 50\n"
        )

    logger.info("Loading indexing config from config.yaml")
    return _parse_index_config(indexing_config)


def _parse_index_config(config_data: dict) -> IndexConfig:
    """Parse raw config dict into IndexConfig."""
    # Get defaults section
    defaults = config_data.get("defaults", {})
    default_include = defaults.get("include", DEFAULT_INCLUDE_PATTERNS.copy())
    default_exclude = defaults.get("exclude", DEFAULT_EXCLUDE_PATTERNS.copy())
    default_recursive = defaults.get("recursive", True)
    default_case_sensitive = defaults.get("case_sensitive", False)

    # Parse directories
    directories: List[DirectoryConfig] = []
    raw_dirs = config_data.get("directories", [])

    if not raw_dirs:
        raise ValueError("Config must have at least one directory in 'directories' list")

    for entry in raw_dirs:
        if isinstance(entry, str):
            # Simple path string - use defaults
            dir_config = DirectoryConfig(
                path=Path(entry),
                include=default_include.copy(),
                exclude=default_exclude.copy(),
                recursive=default_recursive,
                case_sensitive=default_case_sensitive,
            )
        elif isinstance(entry, dict):
            # Full directory config object
            path_str = entry.get("path")
            if not path_str:
                raise ValueError(f"Directory config missing 'path': {entry}")

            dir_config = DirectoryConfig(
                path=Path(path_str),
                enabled=entry.get("enabled", True),
                include=entry.get("include", default_include.copy()),
                exclude=entry.get("exclude", default_exclude.copy()),
                recursive=entry.get("recursive", default_recursive),
                case_sensitive=entry.get("case_sensitive", default_case_sensitive),
            )
        else:
            raise ValueError(f"Invalid directory entry: {entry}")

        directories.append(dir_config)

    return IndexConfig(
        directories=directories,
        chunk_size=config_data.get("chunk_size", 512),
        chunk_overlap=config_data.get("chunk_overlap", 50),
    )


class HeadingStore:
    """Stores document headings separately from chunk metadata.

    This avoids the LlamaIndex SentenceSplitter metadata size validation issue,
    which checks metadata length before applying exclusions. By storing headings
    in a separate file, we keep chunk metadata small while preserving heading
    data for retrieval.

    Supports deferred writes via flush() or context manager to batch disk I/O.
    """

    def __init__(self, store_path: Path):
        self.store_path = store_path
        self._data: Dict[str, List[dict]] = {}
        self._dirty = False  # Track if data needs saving
        self._lock = threading.Lock()  # Thread safety for parallel file loading
        self._load()

    def _load(self):
        """Load heading data from disk."""
        if self.store_path.exists():
            try:
                with open(self.store_path, "r", encoding="utf-8") as f:
                    self._data = json.load(f)
            except Exception as e:
                logger.warning(f"Failed to load heading store: {e}")
                self._data = {}

    def _save(self):
        """Save heading data to disk."""
        self.store_path.parent.mkdir(parents=True, exist_ok=True)
        with open(self.store_path, "w", encoding="utf-8") as f:
            json.dump(self._data, f)

    def set_headings(self, file_path: str, headings: List[dict]):
        """Store headings for a file (deferred write - call flush() to persist)."""
        with self._lock:
            self._data[file_path] = headings
            self._dirty = True

    def get_headings(self, file_path: str) -> List[dict]:
        """Get headings for a file."""
        with self._lock:
            return self._data.get(file_path, [])

    def remove_headings(self, file_path: str):
        """Remove headings for a file (deferred write - call flush() to persist)."""
        with self._lock:
            if file_path in self._data:
                del self._data[file_path]
                self._dirty = True

    def flush(self):
        """Write pending changes to disk."""
        if self._dirty:
            self._save()
            self._dirty = False

    def __enter__(self):
        """Context manager entry."""
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        """Context manager exit - auto-flush on exit."""
        self.flush()
        return False  # Don't suppress exceptions


# Module-level heading store instance (lazy initialized)
_heading_store: Optional["HeadingStore"] = None


def get_heading_store() -> HeadingStore:
    """Get the singleton HeadingStore instance."""
    global _heading_store
    if _heading_store is None:
        _heading_store = HeadingStore(HEADING_STORE_PATH)
    return _heading_store


@dataclass
class FileInfo:
    """Metadata about a file in the data source."""
    path: str
    hash: str
    last_modified: float
    source_dir: str = ""  # Tracks which configured directory this file came from


class IngestionState:
    """Manages the state of ingested files using a SQLite database."""

    def __init__(self, db_path: Path):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        """Initialize the SQLite database schema with migration support."""
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        with sqlite3.connect(self.db_path) as conn:
            # Check if table exists
            cursor = conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name='files'"
            )
            table_exists = cursor.fetchone() is not None

            if not table_exists:
                # Create new table with source_dir column
                conn.execute(
                    """
                    CREATE TABLE files (
                        path TEXT PRIMARY KEY,
                        hash TEXT NOT NULL,
                        last_modified REAL NOT NULL,
                        doc_ids TEXT NOT NULL,
                        source_dir TEXT DEFAULT ''
                    )
                    """
                )
            else:
                # Migration: add source_dir column if missing
                cursor = conn.execute("PRAGMA table_info(files)")
                columns = {row[1] for row in cursor}
                if "source_dir" not in columns:
                    conn.execute("ALTER TABLE files ADD COLUMN source_dir TEXT DEFAULT ''")
                    logger.info("Migrated files table: added source_dir column")

    def get_all_files(self) -> Dict[str, dict]:
        """Retrieve all tracked files and their metadata."""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute(
                "SELECT path, hash, last_modified, doc_ids, source_dir FROM files"
            )
            return {
                row[0]: {
                    "hash": row[1],
                    "last_modified": row[2],
                    "doc_ids": row[3].split(",") if row[3] else [],
                    "source_dir": row[4] if row[4] else "",
                }
                for row in cursor
            }

    def update_file_state(self, file_info: FileInfo, doc_ids: List[str]):
        """Update or insert the state for a file."""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute(
                """
                INSERT INTO files (path, hash, last_modified, doc_ids, source_dir)
                VALUES (?, ?, ?, ?, ?)
                ON CONFLICT(path) DO UPDATE SET
                    hash=excluded.hash,
                    last_modified=excluded.last_modified,
                    doc_ids=excluded.doc_ids,
                    source_dir=excluded.source_dir
                """,
                (
                    file_info.path,
                    file_info.hash,
                    file_info.last_modified,
                    ",".join(doc_ids),
                    file_info.source_dir,
                ),
            )

    def update_file_states_batch(self, file_updates: List[tuple[FileInfo, List[str]]]):
        """Update multiple files in a single transaction (batch operation).

        Args:
            file_updates: List of (FileInfo, doc_ids) tuples
        """
        if not file_updates:
            return

        with sqlite3.connect(self.db_path) as conn:
            conn.executemany(
                """
                INSERT INTO files (path, hash, last_modified, doc_ids, source_dir)
                VALUES (?, ?, ?, ?, ?)
                ON CONFLICT(path) DO UPDATE SET
                    hash=excluded.hash,
                    last_modified=excluded.last_modified,
                    doc_ids=excluded.doc_ids,
                    source_dir=excluded.source_dir
                """,
                [
                    (
                        info.path,
                        info.hash,
                        info.last_modified,
                        ",".join(doc_ids),
                        info.source_dir,
                    )
                    for info, doc_ids in file_updates
                ]
            )

    def remove_file_state(self, path: str):
        """Remove a file from the state tracking."""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("DELETE FROM files WHERE path = ?", (path,))


class DataSource(ABC):
    """Abstract base class for data sources."""

    @abstractmethod
    def iter_files(self, tracked_files: Dict[str, dict] | None = None) -> Iterator[FileInfo]:
        """Yield FileInfo for each file in the source.

        Args:
            tracked_files: Optional dict of previously tracked file states,
                keyed by absolute path. Used for mtime-based fast pre-check.
        """
        pass

    @abstractmethod
    def load_file(
        self,
        file_info: FileInfo,
        ctx: "FileProcessingContext | None" = None
    ) -> List[LlamaIndexDocument]:
        """Load and return documents for a given file.

        Args:
            file_info: File information
            ctx: Optional processing context for progress updates and timeout
        """
        pass


def _compute_line_offsets(text: str) -> List[int]:
    """Compute character offset positions for each line start.

    Returns a list where line_offsets[i] is the character position where line i+1 starts.
    Line 1 starts at position 0 (implicit).
    """
    offsets = [0]  # Line 1 starts at position 0
    for i, char in enumerate(text):
        if char == '\n':
            offsets.append(i + 1)  # Next line starts after the newline
    return offsets


def _extract_markdown_headings(text: str) -> List[dict]:
    """Extract heading hierarchy from Markdown text using ATX-style syntax.

    Parses # Heading syntax and returns list of dicts with text, position, level.
    Handles ATX-style headings (# Heading) but not Setext (underlined).

    Returns:
        List of dicts with keys: text (str), position (int), level (int)
    """
    import re

    headings = []
    # Match ATX-style headings: line start, 1-6 #s, space, text
    pattern = re.compile(r'^(#{1,6})\s+(.+?)$', re.MULTILINE)

    # Find all code block ranges to skip headings inside them
    code_blocks = []
    for match in re.finditer(r'```.*?```', text, flags=re.DOTALL):
        code_blocks.append((match.start(), match.end()))

    def is_in_code_block(pos):
        """Check if position is inside a code block."""
        return any(start <= pos < end for start, end in code_blocks)

    for match in pattern.finditer(text):
        # Skip headings inside code blocks
        if is_in_code_block(match.start()):
            continue

        level = len(match.group(1))
        heading_text = match.group(2).strip()
        position = match.start()

        if heading_text:
            headings.append({
                "text": heading_text,
                "position": position,
                "level": level
            })

    return headings


def _extract_pdf_headings_from_outline(
    pdf_path: Path,
    timeout_event: threading.Event | None = None,
    max_seconds: float = 60.0,
) -> List[dict]:
    """Extract headings from PDF outline/bookmarks (TOC).

    Returns list of dicts with text, position (estimated), level.
    Position is approximate based on cumulative page character counts.
    Falls back to empty list if PDF has no outline or extraction fails.

    Args:
        pdf_path: Path to the PDF file
        timeout_event: Optional threading.Event set when processing should stop
        max_seconds: Wall-clock safety net in seconds (default 60)

    Returns:
        List of dicts with keys: text (str), position (int), level (int).
        Returns partial results on timeout.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("pypdf not available, skipping PDF heading extraction")
        return []

    try:
        reader = PdfReader(pdf_path)
        outline = reader.outline

        if not outline:
            return []

        def flatten_outline(items, level=1):
            """Flatten nested outline into list of (title, page_num, level)."""
            results = []
            for item in items:
                if isinstance(item, list):
                    results.extend(flatten_outline(item, level + 1))
                else:
                    page_num = reader.get_destination_page_number(item)
                    results.append((item.title, page_num, level))
            return results

        flat = flatten_outline(outline)
        headings = []
        start_time = time.time()
        # Cache page text lengths to avoid redundant extract_text() calls
        page_text_lengths: Dict[int, int] = {}

        for title, page_num, level in flat:
            # Check for timeout
            if timeout_event is not None and timeout_event.is_set():
                logger.warning(
                    f"PDF heading extraction timed out for {pdf_path}, "
                    f"returning {len(headings)} partial headings"
                )
                return headings
            if time.time() - start_time >= max_seconds:
                logger.warning(
                    f"PDF heading extraction exceeded {max_seconds}s wall-clock "
                    f"for {pdf_path}, returning {len(headings)} partial headings"
                )
                return headings

            # Estimate position by accumulating text from previous pages
            position = 0
            for page_idx in range(page_num):
                if page_idx < len(reader.pages):
                    if page_idx not in page_text_lengths:
                        page_text_lengths[page_idx] = len(
                            reader.pages[page_idx].extract_text() or ""
                        )
                    position += page_text_lengths[page_idx]

            headings.append({
                "text": title.strip(),
                "position": position,
                "level": level
            })

        return headings

    except Exception as e:
        logger.warning(f"Failed to extract PDF outline from {pdf_path}: {e}")
        return []


def _split_docx_with_timeout(
    docx_path: Path,
    ctx: "FileProcessingContext | None",
    timeout_seconds: float,
) -> List[LlamaIndexDocument]:
    """Run split_docx_into_heading_documents() in a worker thread with a hard timeout.

    Returns the loaded documents, or an empty list if the call times out.
    """
    with ThreadPoolExecutor(max_workers=1) as pool:
        future = pool.submit(split_docx_into_heading_documents, docx_path, ctx)
        try:
            return future.result(timeout=timeout_seconds)
        except Exception as e:
            future.cancel()
            if "TimeoutError" in type(e).__name__ or isinstance(e, TimeoutError):
                logger.warning(
                    f"DOCX processing timed out after {timeout_seconds:.0f}s: {docx_path}"
                )
                return []
            raise


def _load_data_with_timeout(
    reader: SimpleDirectoryReader, timeout_seconds: float
) -> List[LlamaIndexDocument]:
    """Run reader.load_data() in a worker thread with a hard timeout.

    Returns the loaded documents, or an empty list if the call times out.
    """
    with ThreadPoolExecutor(max_workers=1) as pool:
        future = pool.submit(reader.load_data)
        try:
            return future.result(timeout=timeout_seconds)
        except Exception as e:
            future.cancel()
            if "TimeoutError" in type(e).__name__ or isinstance(e, TimeoutError):
                logger.warning(
                    f"load_data() timed out after {timeout_seconds:.0f}s"
                )
                return []
            raise


class LocalFileSystemSource(DataSource):
    """Implementation of DataSource for the local file system with filtering."""

    def __init__(self, config: DirectoryConfig):
        self.config = config
        self.base_dir = config.path

    def is_available(self) -> bool:
        """Check if the directory is available and accessible."""
        try:
            if not self.base_dir.exists():
                return False
            if not self.base_dir.is_dir():
                return False
            # Try to list directory to verify access (important for network mounts)
            next(self.base_dir.iterdir(), None)
            return True
        except (OSError, PermissionError):
            return False

    def _matches_patterns(self, file_path: Path) -> bool:
        """Check if file matches include patterns and doesn't match exclude patterns.

        Uses PurePosixPath.match() for glob pattern matching.
        For directory exclusion patterns like **/*venv*/**, checks each path component.
        When case_sensitive is False (default), matching is case-insensitive.
        """
        import fnmatch
        from pathlib import PurePosixPath

        try:
            rel_path = file_path.relative_to(self.base_dir)
        except ValueError:
            rel_path = Path(file_path.name)

        ci = not self.config.case_sensitive
        abs_str = str(file_path).lower() if ci else str(file_path)
        rel_str = str(rel_path).lower() if ci else str(rel_path)
        name = file_path.name.lower() if ci else file_path.name

        # Check exclude patterns first
        for pattern in self.config.exclude:
            pat = pattern.lower() if ci else pattern
            # Handle directory exclusion patterns (e.g., **/*venv*/**, **/node_modules/**)
            # by checking if any directory component matches
            if pattern.startswith('**/') and pattern.endswith('/**'):
                # Extract the directory pattern (e.g., *venv* or node_modules)
                dir_pattern = pat[3:-3]  # Remove **/ prefix and /** suffix
                for part in rel_path.parts[:-1]:  # Check all directory components (not filename)
                    if fnmatch.fnmatch(part.lower() if ci else part, dir_pattern):
                        return False
            else:
                # Standard pattern matching
                if PurePosixPath(rel_str).match(pat) or name == pat:
                    return False

        # Check include patterns
        if not self.config.include:
            return True

        for pattern in self.config.include:
            pat = pattern.lower() if ci else pattern
            if PurePosixPath(rel_str).match(pat) or PurePosixPath(abs_str).match(pat):
                return True

        return False

    def _should_skip_directory(self, dir_name: str) -> bool:
        """Check if a directory should be pruned from os.walk traversal.

        Handles patterns of the form **/<pattern>/** by checking the directory name.
        """
        import fnmatch

        for pattern in self.config.exclude:
            if pattern.startswith('**/') and pattern.endswith('/**'):
                dir_pattern = pattern[3:-3]  # e.g., "node_modules", ".git", "*venv*"
                if fnmatch.fnmatch(dir_name, dir_pattern):
                    return True
        return False

    def iter_files(self, tracked_files: Dict[str, dict] | None = None) -> Iterator[FileInfo]:
        """Yield FileInfo for each matching file in the source.

        Args:
            tracked_files: Optional dict of previously tracked file states,
                keyed by absolute path. Used for mtime-based fast pre-check.
        """
        if self.config.recursive:
            # topdown=True (default) allows pruning dirs in-place
            for root, dirs, files in os.walk(self.base_dir):
                # Prune excluded directories in-place to prevent descent
                dirs[:] = [d for d in dirs if not self._should_skip_directory(d)]

                root_path = Path(root)
                for file in files:
                    file_path = root_path / file
                    if not self._matches_patterns(file_path):
                        continue
                    try:
                        yield self._create_file_info(file_path, tracked_files)
                    except (OSError, IOError) as e:
                        logger.warning(f"Could not access file {file_path}: {e}")
                        continue
        else:
            # Non-recursive: only top-level files
            try:
                top_files = [
                    f for f in self.base_dir.iterdir() if f.is_file()
                ]
            except OSError as e:
                logger.warning(f"Could not list directory {self.base_dir}: {e}")
                return

            for f in top_files:
                if not self._matches_patterns(f):
                    continue
                try:
                    yield self._create_file_info(f, tracked_files)
                except (OSError, IOError) as e:
                    logger.warning(f"Could not access file {f}: {e}")
                    continue

    def _create_file_info(
        self,
        file_path: Path,
        tracked_files: Dict[str, dict] | None = None,
    ) -> FileInfo:
        """Create FileInfo with source directory context.

        If tracked_files is provided and the file's mtime is unchanged,
        reuses the cached hash to avoid reading the entire file.
        """
        abs_path = str(file_path.absolute())
        mtime = file_path.stat().st_mtime

        # Fast path: reuse cached hash when mtime is unchanged
        if tracked_files is not None:
            cached = tracked_files.get(abs_path)
            if cached is not None and cached.get("last_modified") == mtime:
                return FileInfo(
                    path=abs_path,
                    hash=cached["hash"],
                    last_modified=mtime,
                    source_dir=str(self.base_dir.absolute()),
                )

        # Slow path: compute MD5
        hasher = hashlib.md5()
        with open(file_path, "rb") as f:
            buf = f.read(65536)
            while len(buf) > 0:
                hasher.update(buf)
                buf = f.read(65536)

        return FileInfo(
            path=abs_path,
            hash=hasher.hexdigest(),
            last_modified=mtime,
            source_dir=str(self.base_dir.absolute()),
        )

    def load_file(
        self,
        file_info: FileInfo,
        ctx: "FileProcessingContext | None" = None
    ) -> List[LlamaIndexDocument]:
        file_path = Path(file_info.path)
        if not file_path.exists():
            logger.warning(f"Skipping disappeared file: {file_path}")
            return []
        if file_path.suffix.lower() == ".docx":
            if ctx:
                ctx.set_phase("Parsing DOCX")
            remaining = ctx.remaining_seconds() if ctx else None
            if remaining is not None:
                return _split_docx_with_timeout(file_path, ctx, remaining)
            return split_docx_into_heading_documents(file_path, ctx)
        elif file_path.suffix.lower() == ".doc":
            # Convert .doc to .docx using LibreOffice, then process
            if ctx:
                ctx.set_phase("Converting .doc to .docx")

            # Use specialized timeout for .doc conversion
            doc_timeout = cfgload.get("indexing.timeout.doc_conversion_seconds", 90)
            docx_path = _convert_doc_to_docx(file_path, timeout=doc_timeout)

            if docx_path is None:
                logger.warning(f"Skipping {file_path}: could not convert .doc to .docx")
                return []
            try:
                if ctx:
                    ctx.set_phase("Parsing converted DOCX")
                remaining = ctx.remaining_seconds() if ctx else None
                if remaining is not None:
                    docs = _split_docx_with_timeout(docx_path, ctx, remaining)
                else:
                    docs = split_docx_into_heading_documents(docx_path, ctx)
                # Update metadata to point to original .doc file
                for doc in docs:
                    doc.metadata["file_path"] = str(file_path)
                    doc.metadata["file_name"] = file_path.name
                    if "source" in doc.metadata:
                        doc.metadata["source"] = str(file_path)
                return docs
            finally:
                # Clean up temp file
                if docx_path.exists():
                    docx_path.unlink()
        else:
            reader = SimpleDirectoryReader(
                input_files=[str(file_path)],
            )
            remaining = ctx.remaining_seconds() if ctx else None
            if remaining is not None:
                docs = _load_data_with_timeout(reader, remaining)
            else:
                docs = reader.load_data()
            # Ensure dates are visible to LLM (remove from exclusion list)
            for doc in docs:
                if hasattr(doc, 'excluded_llm_metadata_keys') and doc.excluded_llm_metadata_keys:
                    doc.excluded_llm_metadata_keys = [
                        k for k in doc.excluded_llm_metadata_keys
                        if k not in ('creation_date', 'last_modified_date')
                    ]

            # Add line offsets for text-based files (markdown, txt) to enable line number lookup
            if file_path.suffix.lower() in {".md", ".txt"}:
                for doc in docs:
                    text = doc.get_content()
                    line_offsets = _compute_line_offsets(text)
                    doc.metadata["line_offsets"] = line_offsets

                    # Extract headings for Markdown and store separately
                    # (not in metadata to avoid SentenceSplitter size validation)
                    if file_path.suffix.lower() == ".md":
                        headings = _extract_markdown_headings(text)
                        get_heading_store().set_headings(str(file_path), headings)

            # Extract headings for PDF files and store separately
            if file_path.suffix.lower() == ".pdf":
                timeout_event = ctx._timeout_event if ctx else None
                headings = _extract_pdf_headings_from_outline(
                    file_path, timeout_event=timeout_event
                )
                get_heading_store().set_headings(str(file_path), headings)
                if ctx:
                    ctx.check_timeout()

            # Apply metadata exclusions
            for doc in docs:
                doc.excluded_embed_metadata_keys = EXCLUDED_EMBED_METADATA_KEYS
                doc.excluded_llm_metadata_keys = EXCLUDED_LLM_METADATA_KEYS

            return docs


class MultiDirectoryDataSource(DataSource):
    """Aggregates multiple LocalFileSystemSource instances."""

    def __init__(self, config: IndexConfig):
        self.config = config
        self.sources: List[LocalFileSystemSource] = []
        self.unavailable_dirs: List[DirectoryConfig] = []

        for dir_config in config.directories:
            if not dir_config.enabled:
                logger.info(f"Skipping disabled directory: {dir_config.path}")
                continue

            source = LocalFileSystemSource(dir_config)

            if source.is_available():
                self.sources.append(source)
                logger.info(f"Added directory source: {dir_config.path}")
            else:
                self.unavailable_dirs.append(dir_config)
                logger.warning(f"Directory unavailable, skipping: {dir_config.path}")

    def iter_files(self, tracked_files: Dict[str, dict] | None = None) -> Iterator[FileInfo]:
        """Iterate over files from all available sources."""
        seen_paths: Set[str] = set()

        for source in self.sources:
            for file_info in source.iter_files(tracked_files=tracked_files):
                # Deduplicate in case of overlapping mounts
                if file_info.path not in seen_paths:
                    seen_paths.add(file_info.path)
                    yield file_info

    def load_file(
        self,
        file_info: FileInfo,
        ctx: "FileProcessingContext | None" = None
    ) -> List[LlamaIndexDocument]:
        """Load file using the appropriate source based on source_dir."""
        # Find the source that owns this file
        for source in self.sources:
            if file_info.source_dir == str(source.base_dir.absolute()):
                return source.load_file(file_info, ctx)

        # Fallback: use first source (shouldn't happen normally)
        if self.sources:
            return self.sources[0].load_file(file_info, ctx)

        raise ValueError(f"No source available for file: {file_info.path}")

    def get_summary(self) -> Dict[str, Any]:
        """Return summary of configured directories."""
        return {
            "available": [str(s.base_dir) for s in self.sources],
            "unavailable": [str(d.path) for d in self.unavailable_dirs],
            "total_sources": len(self.sources),
        }


class _DevNull:
    """Minimal file-like sink that discards all writes."""

    def write(self, _data: str) -> int:
        return 0

    def flush(self) -> None:
        pass

    def isatty(self) -> bool:
        return False


class IndexingUI:
    """Unified terminal output for the indexing pipeline.

    Owns all stdout writes during build_index(). Provides two display modes:
    - Step mode: "message... done" with animated spinner
    - Progress mode: progress bar with optional sub-line for current file

    All methods are thread-safe via a single lock.
    """

    _SPINNER_CHARS = "⠋⠙⠹⠸⠼⠴⠦⠧⠇⠏"

    def __init__(self, stream=None, verbose=False):
        self._stream = stream or sys.stdout
        self._verbose = verbose
        self._lock = threading.Lock()

        # ANSI color codes (disabled when stream is not a TTY)
        _s = self._stream
        self._tty = hasattr(_s, "isatty") and _s.isatty()
        self.RESET = "\033[0m" if self._tty else ""
        self.BOLD = "\033[1m" if self._tty else ""
        self.DIM = "\033[2m" if self._tty else ""
        self.GREEN = "\033[32m" if self._tty else ""
        self.YELLOW = "\033[33m" if self._tty else ""
        self.CYAN = "\033[36m" if self._tty else ""
        self.RED = "\033[31m" if self._tty else ""
        self.BOLD_GREEN = "\033[1;32m" if self._tty else ""
        self.BOLD_CYAN = "\033[1;36m" if self._tty else ""

        # Step/spinner state
        self._step_message: str | None = None
        self._step_original_message: str | None = None
        self._step_stop = threading.Event()
        self._step_thread: threading.Thread | None = None

        # Progress bar state
        self._progress_active = False
        self._progress_paused = False
        self._progress_total = 0
        self._progress_current = 0
        self._progress_desc = ""
        self._progress_unit = "file"
        self._progress_width = 30
        self._progress_file = ""
        self._progress_phase = ""
        self._progress_heartbeat = ""
        self._progress_has_subline = False
        self._progress_last_pct = -1  # last printed percentage (non-TTY)
        self._progress_substep = ""
        self._progress_substep_stop = threading.Event()
        self._progress_substep_thread: threading.Thread | None = None
        self._progress_substep_idx = 0

        # Output suppression state (populated by _suppress_output)
        self._orig_stdout = None
        self._orig_stderr = None
        self._orig_handler_levels: list[tuple[logging.Handler, int]] = []

    # -- context manager --

    def __enter__(self):
        self._suppress_output()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        if self._step_message is not None:
            self.step_done("interrupted")
        self._step_original_message = None
        if self._progress_substep:
            self.progress_substep_done()
        if self._progress_active or self._progress_paused:
            self._progress_active = True
            self._progress_paused = False
            self.progress_done()
        self._restore_output()
        return False

    # -- step mode --

    def step_start(self, message: str) -> None:
        """Begin an animated step: prints 'message... ⠋' with spinner."""
        with self._lock:
            self._step_message = message
            self._step_original_message = message
            self._step_stop.clear()
            if not self._tty:
                self._write(f"{message}... ")
                return
            self._write(f"\r{self.BOLD}{message}{self.RESET}... ")
        self._step_thread = threading.Thread(target=self._step_spin, daemon=True)
        self._step_thread.start()

    def step_update(self, message: str) -> None:
        """Update the step message while the spinner continues."""
        with self._lock:
            if self._step_message is not None:
                self._step_message = message

    def step_done(self, suffix: str = "done") -> None:
        """Complete the current step: replaces spinner with suffix.

        Uses the original step_start message for the final line so that
        dynamic updates (e.g. file counts) don't appear in the completed line.
        """
        self._step_stop.set()
        if self._step_thread:
            self._step_thread.join()
            self._step_thread = None
        with self._lock:
            msg = self._step_original_message or self._step_message or ""
            self._step_message = None
            self._step_original_message = None
            if not self._tty:
                self._write(f"{suffix}\n")
                return
            colored_suffix = self._color_suffix(suffix)
            self._write(f"\r{self.BOLD}{msg}{self.RESET}... {colored_suffix}\033[K\n")

    def _step_spin(self) -> None:
        idx = 0
        while not self._step_stop.is_set():
            with self._lock:
                if self._step_message is not None:
                    self._write(
                        f"\r{self.BOLD}{self._step_message}{self.RESET}... "
                        f"{self.CYAN}{self._SPINNER_CHARS[idx]}{self.RESET}\033[K"
                    )
            idx = (idx + 1) % len(self._SPINNER_CHARS)
            self._step_stop.wait(0.1)

    # -- progress mode --

    def progress_start(self, total: int, desc: str = "Processing files", unit: str = "file") -> None:
        """Enter progress bar mode."""
        with self._lock:
            self._progress_active = True
            self._progress_paused = False
            self._progress_total = max(total, 0)
            self._progress_current = 0
            self._progress_desc = desc
            self._progress_unit = unit
            self._progress_file = ""
            self._progress_phase = ""
            self._progress_heartbeat = ""
            self._progress_has_subline = False
            self._progress_last_pct = -1
            if self._progress_total > 0:
                self._render_progress()

    def progress_update(self, step: int = 1) -> None:
        """Advance the progress bar."""
        with self._lock:
            if not self._progress_active or self._progress_total <= 0:
                return
            self._progress_current = min(self._progress_total, self._progress_current + step)
            self._render_progress()

    def progress_set_file(self, file_path: str, phase: str = "") -> None:
        """Set current file shown on sub-line under the bar."""
        with self._lock:
            if not self._progress_active:
                return
            self._progress_file = file_path
            self._progress_phase = phase
            self._render_progress()

    def progress_set_heartbeat(self, char: str) -> None:
        """Update heartbeat animation character."""
        with self._lock:
            if not self._progress_active:
                return
            self._progress_heartbeat = char
            self._render_progress()

    def progress_pause(self) -> None:
        """Temporarily hide the progress bar for step output."""
        with self._lock:
            if not self._progress_active:
                return
            self._clear_progress_area()
            self._progress_active = False
            self._progress_paused = True

    def progress_resume(self) -> None:
        """Re-show the progress bar after a pause."""
        with self._lock:
            if not self._progress_paused:
                return
            self._progress_active = True
            self._progress_paused = False
            self._progress_has_subline = False
            self._render_progress()

    def progress_substep_start(self, message: str) -> None:
        """Show a substep message on the progress bar sub-line with spinner."""
        with self._lock:
            self._progress_substep = message
            self._progress_file = ""
            self._progress_phase = ""
            self._progress_heartbeat = ""
            self._progress_substep_stop.clear()
            self._progress_substep_idx = 0
            if not self._tty:
                # Print current progress context + substep on one line
                pct_int = 0
                if self._progress_total > 0:
                    pct_int = int(self._progress_current / self._progress_total * 100)
                self._stream.write(
                    f"{self._progress_desc} "
                    f"[{self._progress_current}/{self._progress_total}] "
                    f"{pct_int}% \u2014 {message}... "
                )
                self._stream.flush()
                return
            self._render_progress()
        self._progress_substep_thread = threading.Thread(
            target=self._substep_spin, daemon=True
        )
        self._progress_substep_thread.start()

    def progress_substep_done(self) -> None:
        """Clear the substep message from the progress bar sub-line."""
        self._progress_substep_stop.set()
        if self._progress_substep_thread:
            self._progress_substep_thread.join()
            self._progress_substep_thread = None
        with self._lock:
            self._progress_substep = ""
            self._progress_heartbeat = ""
            if not self._tty:
                self._stream.write("done\n")
                self._stream.flush()
                return
            self._render_progress()

    def _substep_spin(self) -> None:
        """Animate spinner on the progress sub-line for substep."""
        while not self._progress_substep_stop.is_set():
            with self._lock:
                if self._progress_substep:
                    self._progress_heartbeat = self._SPINNER_CHARS[self._progress_substep_idx]
                    self._render_progress()
            self._progress_substep_idx = (self._progress_substep_idx + 1) % len(self._SPINNER_CHARS)
            self._progress_substep_stop.wait(0.1)

    def progress_done(self) -> None:
        """Exit progress bar mode."""
        with self._lock:
            if not self._progress_active:
                return
            if not self._tty:
                # Print final 100% line if not already printed
                if self._progress_last_pct < 100 and self._progress_total > 0:
                    self._write(
                        f"{self._progress_desc} "
                        f"[{self._progress_total}/{self._progress_total}] 100%\n"
                    )
            else:
                # Render final state
                self._render_progress()
                # Move past the progress area
                if self._progress_has_subline:
                    self._write("\n\n")
                else:
                    self._write("\n")
            self._progress_active = False
            self._progress_paused = False
            self._progress_has_subline = False

    def _render_progress(self) -> None:
        """Render progress bar + optional file sub-line. Must hold lock."""
        if self._progress_total <= 0:
            return
        progress = self._progress_current / self._progress_total

        if not self._tty:
            # Non-TTY: print a simple line every 10% to avoid log spam
            pct_int = int(progress * 100)
            threshold = (pct_int // 10) * 10
            if threshold <= self._progress_last_pct:
                return
            self._progress_last_pct = threshold
            self._stream.write(
                f"{self._progress_desc} "
                f"[{self._progress_current}/{self._progress_total}] "
                f"{pct_int}%\n"
            )
            self._stream.flush()
            return

        filled = int(self._progress_width * progress)
        bar_filled = f"{self.GREEN}{'█' * filled}{self.RESET}"
        bar_empty = f"{self.DIM}{'░' * (self._progress_width - filled)}{self.RESET}"
        bar = f"{bar_filled}{bar_empty}"

        # Move cursor to start of progress area
        if self._progress_has_subline:
            self._stream.write("\033[1A\r")

        # Line 1: progress bar
        pct = f"{progress * 100:5.1f}%"
        if progress >= 1.0:
            pct = f"{self.BOLD_GREEN}{pct}{self.RESET}"
        line1 = (
            f"{self.BOLD}{self._progress_desc}{self.RESET} [{bar}] "
            f"{pct}  {self.DIM}({self._progress_current}/{self._progress_total}){self.RESET}"
        )
        self._stream.write(f"\r\033[K{line1}")

        # Line 2: substep message (priority) or current file
        if self._progress_substep:
            subline = f"  {self.DIM}{self._progress_substep}{self.RESET}"
            if self._progress_heartbeat:
                subline += f" {self.CYAN}{self._progress_heartbeat}{self.RESET}"
            self._stream.write(f"\n\033[K{subline}")
            self._progress_has_subline = True
        elif self._progress_file:
            file_display = Path(self._progress_file).name
            if len(file_display) > 50:
                file_display = "..." + file_display[-47:]
            subline = f"  {self.DIM}{file_display}{self.RESET}"
            if self._progress_phase:
                subline += f" {self.DIM}({self._progress_phase}"
                if self._progress_heartbeat:
                    subline += f" {self.CYAN}{self._progress_heartbeat}{self.RESET}{self.DIM}"
                subline += f"){self.RESET}"
            elif self._progress_heartbeat:
                subline += f" {self.CYAN}{self._progress_heartbeat}{self.RESET}"
            self._stream.write(f"\n\033[K{subline}")
            self._progress_has_subline = True
        elif self._progress_has_subline:
            # Clear stale sub-line
            self._stream.write(f"\n\033[K")

        self._stream.flush()

    def _clear_progress_area(self) -> None:
        """Clear progress bar lines from terminal. Must hold lock."""
        if not self._tty:
            return
        if self._progress_has_subline:
            self._stream.write("\033[1A\r\033[K\n\033[K\033[1A\r")
        else:
            self._stream.write("\r\033[K")
        self._stream.flush()
        self._progress_has_subline = False

    # -- general output --

    def print(self, message: str) -> None:
        """Print a plain text line."""
        with self._lock:
            self._write(f"{message}\n")

    def success(self, message: str) -> None:
        """Print a success message in bold green."""
        with self._lock:
            self._write(f"{self.BOLD_GREEN}{message}{self.RESET}\n")

    def error(self, message: str) -> None:
        """Print an error message in red."""
        with self._lock:
            self._write(f"{self.RED}{message}{self.RESET}\n")

    # -- internal helpers --

    def _color_suffix(self, suffix: str) -> str:
        """Return a color-coded suffix string for step_done output."""
        s = suffix.lower()
        if s in ("done", "no changes"):
            return f"{self.GREEN}{suffix}{self.RESET}"
        if s in ("skipped",):
            return f"{self.YELLOW}{suffix}{self.RESET}"
        if s in ("interrupted",):
            return f"{self.RED}{suffix}{self.RESET}"
        return suffix

    def _write(self, text: str) -> None:
        """Write to stream and flush. Caller must hold lock."""
        self._stream.write(text)
        self._stream.flush()

    def _suppress_output(self) -> None:
        """Redirect stdout/stderr and silence root logger stream handlers.

        IndexingUI captures self._stream at __init__ time, so it keeps
        writing to the real terminal. All 3rd-party code that calls
        print() or writes to sys.stdout/stderr hits _DevNull instead.

        Skipped when verbose=True to allow full debugging output.
        """
        if self._verbose:
            return

        devnull = _DevNull()

        self._orig_stdout = sys.stdout
        self._orig_stderr = sys.stderr
        sys.stdout = devnull
        sys.stderr = devnull

        for handler in logging.root.handlers:
            if isinstance(handler, logging.StreamHandler) and not isinstance(handler, logging.FileHandler):
                self._orig_handler_levels.append((handler, handler.level))
                handler.setLevel(logging.CRITICAL + 1)

    def _restore_output(self) -> None:
        """Restore stdout/stderr and root logger handler levels."""
        if self._orig_stdout is not None:
            sys.stdout = self._orig_stdout
            self._orig_stdout = None
        if self._orig_stderr is not None:
            sys.stderr = self._orig_stderr
            self._orig_stderr = None

        for handler, level in self._orig_handler_levels:
            handler.setLevel(level)
        self._orig_handler_levels.clear()


class FileProcessingTimeoutError(Exception):
    """Raised when file processing exceeds timeout."""
    pass


class FileProcessingContext:
    """Context manager for file processing with timeout and heartbeat.

    Usage:
        with FileProcessingContext(file_path, ui, timeout=300) as ctx:
            ctx.set_phase("Converting .doc")
            result = process_file()
    """

    def __init__(
        self,
        file_path: str,
        ui: IndexingUI,
        timeout_seconds: float | None = None,
        heartbeat_interval: float = 2.0
    ):
        self.file_path = file_path
        self.ui = ui
        self.timeout_seconds = timeout_seconds
        self.heartbeat_interval = heartbeat_interval

        self._start_time: float | None = None
        self._stop_event = threading.Event()
        self._timeout_event = threading.Event()
        self._heartbeat_thread: threading.Thread | None = None
        self._current_phase = ""

    def __enter__(self):
        """Start timing and heartbeat thread."""
        self._start_time = time.time()

        # Start heartbeat thread
        self._heartbeat_thread = threading.Thread(
            target=self._heartbeat_loop,
            daemon=True
        )
        self._heartbeat_thread.start()

        # Update UI with current file
        self.ui.progress_set_file(self.file_path, "")

        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        """Stop heartbeat and check for slow files."""
        # Stop heartbeat thread
        self._stop_event.set()
        if self._heartbeat_thread:
            self._heartbeat_thread.join(timeout=1.0)

        # Warn about slow files (warnings still pass through)
        if self._start_time:
            duration = time.time() - self._start_time
            slow_threshold = cfgload.get(
                "indexing.logging.slow_file_threshold_seconds", 30
            )
            if duration > slow_threshold:
                logger.warning(
                    f"Slow file processing: {self.file_path} took {duration:.1f}s"
                )

        # Don't suppress exceptions
        return False

    def set_phase(self, phase: str) -> None:
        """Update current operation phase."""
        self._current_phase = phase
        self.ui.progress_set_file(self.file_path, phase)

        # Check for timeout
        self.check_timeout()

    def check_timeout(self) -> None:
        """Check if processing has exceeded timeout (main-thread safe).

        Checks both the _timeout_event (set by heartbeat thread) and
        elapsed wall-clock time. Raises FileProcessingTimeoutError if
        either indicates a timeout.
        """
        if self.timeout_seconds is None or self._start_time is None:
            return

        if self._timeout_event.is_set():
            elapsed = time.time() - self._start_time
            raise FileProcessingTimeoutError(
                f"File processing timed out after {elapsed:.1f}s: {self.file_path}"
            )

        elapsed = time.time() - self._start_time
        if elapsed > self.timeout_seconds:
            self._timeout_event.set()
            raise FileProcessingTimeoutError(
                f"File processing timed out after {elapsed:.1f}s: {self.file_path}"
            )

    def remaining_seconds(self) -> float | None:
        """Return seconds remaining before timeout, or None if no timeout set."""
        if self.timeout_seconds is None or self._start_time is None:
            return None
        remaining = self.timeout_seconds - (time.time() - self._start_time)
        return max(0.0, remaining)

    def _check_timeout(self) -> None:
        """Check if processing has exceeded timeout (used by heartbeat thread)."""
        if self.timeout_seconds is None or self._start_time is None:
            return

        elapsed = time.time() - self._start_time
        if elapsed > self.timeout_seconds:
            self._timeout_event.set()
            raise FileProcessingTimeoutError(
                f"File processing timed out after {elapsed:.1f}s: {self.file_path}"
            )

    def _heartbeat_loop(self) -> None:
        """Background thread that updates heartbeat indicator."""
        if not self.ui._tty:
            # Non-TTY: only monitor for timeouts, skip animation
            while not self._stop_event.is_set():
                try:
                    self._check_timeout()
                except FileProcessingTimeoutError:
                    break
                time.sleep(self.heartbeat_interval)
            return

        spinner_chars = "⠋⠙⠹⠸⠼⠴⠦⠧⠇⠏"
        idx = 0

        while not self._stop_event.is_set():
            self.ui.progress_set_heartbeat(spinner_chars[idx])
            idx = (idx + 1) % len(spinner_chars)

            # Check for timeout
            try:
                self._check_timeout()
            except FileProcessingTimeoutError:
                break

            time.sleep(self.heartbeat_interval)


def _embedding_cache_path(model_name: str, cache_dir: Path) -> Path:
    """Return the expected cache directory for a FastEmbed model."""
    return cache_dir / f"models--{model_name.replace('/', '--')}"


def _verify_model_cache_exists(cache_dir: Path) -> bool:
    """
    Verify that the cached model directory exists and contains the expected model files.
    """
    from fastembed import TextEmbedding

    try:
        models = TextEmbedding.list_supported_models()
        model_info = [m for m in models if m.get("model") == RETRIEVAL_EMBED_MODEL_NAME]
        if not model_info:
            return False

        model_info = model_info[0]
        hf_source = model_info.get("sources", {}).get("hf")
        if not hf_source:
            return False

        expected_dir = cache_dir / f"models--{hf_source.replace('/', '--')}"
        if not expected_dir.exists():
            return False

        snapshots_dir = expected_dir / "snapshots"
        if not snapshots_dir.exists():
            return False

        model_file = model_info.get("model_file", "model_optimized.onnx")
        for snapshot in snapshots_dir.iterdir():
            if snapshot.is_dir():
                model_path = snapshot / model_file
                if model_path.exists() or model_path.is_symlink():
                    return True

        return False
    except Exception:
        return False


def _get_cached_model_path(cache_dir: Path, model_name: str) -> Path | None:
    """Get the cached model directory path."""
    try:
        from huggingface_hub import snapshot_download
        from fastembed import TextEmbedding
        models = TextEmbedding.list_supported_models()
        model_info = [m for m in models if m.get("model") == model_name]
        if model_info:
            hf_source = model_info[0].get("sources", {}).get("hf")
            if hf_source:
                cache_dir_abs = cache_dir.resolve()
                model_dir = snapshot_download(
                    repo_id=hf_source,
                    local_files_only=True,
                    cache_dir=str(cache_dir_abs)
                )
                return Path(model_dir).resolve()
    except (ImportError, Exception):
        pass
    return None


def _create_fastembed_embedding(cache_dir: Path, offline: bool = False):
    """Create a FastEmbedEmbedding instance."""
    # Large batch size so LlamaIndex sends all texts in one call to
    # _get_text_embeddings(), avoiding per-batch instrumentation overhead
    # (to_dict serialization, event dispatch). Default is 10, which causes
    # ~38 separate cycles for 375 chunks. FastEmbed handles its own internal
    # ONNX batching (default 256).
    embed_batch_size = 512
    threads = os.cpu_count()

    if offline:
        cached_model_path = _get_cached_model_path(cache_dir, RETRIEVAL_EMBED_MODEL_NAME)
        if cached_model_path:
            logger.info(
                f"Using cached model path to bypass download: {cached_model_path}"
            )
            return FastEmbedEmbedding(
                model_name=RETRIEVAL_EMBED_MODEL_NAME,
                cache_dir=str(cache_dir),
                specific_model_path=str(cached_model_path),
                embed_batch_size=embed_batch_size,
                threads=threads,
            )
        else:
            logger.warning(
                "Could not find cached model path, falling back to normal initialization"
            )

    return FastEmbedEmbedding(
        model_name=RETRIEVAL_EMBED_MODEL_NAME,
        cache_dir=str(cache_dir),
        embed_batch_size=embed_batch_size,
        threads=threads,
    )


def ensure_embedding_model_cached(cache_dir: Path, offline: bool = False) -> None:
    """Ensure the embedding model is available in the local cache."""
    if offline:
        logger.info("Verifying embedding model cache...")
        if _verify_model_cache_exists(cache_dir):
            logger.info("Embedding model found in cache")
        else:
            logger.error(
                "Offline mode enabled, but embedding model cache not found in %s",
                cache_dir,
            )
            raise FileNotFoundError(
                f"Embedding model '{RETRIEVAL_EMBED_MODEL_NAME}' not found in cache directory '{cache_dir}'. "
            )

    try:
        logger.info("Initializing embedding model from cache...")
        cache_dir_abs = cache_dir.resolve()
        if offline:
            os.environ["HF_HUB_CACHE"] = str(cache_dir_abs)

        _create_fastembed_embedding(cache_dir, offline=offline)
        logger.info("Embedding model initialized successfully")
        return
    except (ValueError, Exception) as e:
        # Simplified error handling for brevity, similar logic as original
        if offline:
            raise FileNotFoundError(f"Failed to load model offline: {e}") from e
        else:
            raise RuntimeError(f"Failed to download/initialize model: {e}") from e


def ensure_rerank_model_cached(cache_dir: Path, offline: bool = False) -> Path:
    """Ensure the reranking model is cached locally."""
    try:
        from flashrank import Ranker
    except ImportError as exc:
        raise ImportError(
            "flashrank is required for reranking."
        ) from exc

    cache_dir_abs = cache_dir.resolve()
    logger.info("Ensuring rerank model is available in cache...")

    # Map cross-encoder model names to FlashRank equivalents if needed
    model_name = RETRIEVAL_RERANK_MODEL_NAME
    # Note: FlashRank doesn't have L-6 models, so we map to L-12 equivalents
    model_mapping = {
        "cross-encoder/ms-marco-MiniLM-L-6-v2": "ms-marco-MiniLM-L-12-v2",  # L-6 not available, use L-12
        "ms-marco-MiniLM-L-6-v2": "ms-marco-MiniLM-L-12-v2",  # Direct mapping for L-6
    }
    if model_name in model_mapping:
        model_name = model_mapping[model_name]
    elif model_name.startswith("cross-encoder/"):
        # Extract model name after cross-encoder/ prefix and try to map
        base_name = model_name.replace("cross-encoder/", "")
        # If it's an L-6 model, map to L-12
        if "L-6" in base_name:
            model_name = base_name.replace("L-6", "L-12")
        else:
            model_name = base_name

    try:
        reranker = Ranker(model_name=model_name, cache_dir=str(cache_dir_abs))
        logger.info(f"FlashRank model '{model_name}' initialized successfully")
        return cache_dir_abs
    except Exception as exc:
        if offline:
            raise FileNotFoundError(
                f"Rerank model '{model_name}' not found in cache."
            ) from exc
        raise


def _parse_heading_level(style_name: str | None) -> int:
    """Best-effort extraction of a numeric heading level from a DOCX style name."""
    if not style_name:
        return 1
    try:
        if "Heading" in style_name:
            level_str = style_name.replace("Heading", "").strip()
            if level_str:
                return int(level_str)
    except (ValueError, AttributeError):
        pass
    return 1


def _get_doc_temp_dir() -> Path:
    """Get the temporary directory for .doc conversion, creating it if needed."""
    storage_dir = Path(cfgload.get("storage.storage_dir", "./storage"))
    temp_dir = storage_dir / "doc_temp"
    temp_dir.mkdir(parents=True, exist_ok=True)
    return temp_dir


def _convert_doc_to_docx(doc_path: Path, timeout: float = 60) -> Path | None:
    """Convert a .doc file to .docx using LibreOffice.

    Args:
        doc_path: Path to .doc file
        timeout: Timeout in seconds for conversion process

    Returns:
        Path to temporary .docx file, or None if conversion fails.
        Caller is responsible for cleaning up the temp file.
    """
    import subprocess
    import shutil

    # Find LibreOffice executable
    soffice_paths = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
        "/usr/bin/soffice",  # Linux
        "/usr/bin/libreoffice",  # Linux alternative
        "soffice",  # Windows (in PATH)
    ]

    soffice = None
    for path in soffice_paths:
        if shutil.which(path):
            soffice = path
            break

    if not soffice:
        logger.warning(f"LibreOffice not found. Cannot convert {doc_path}")
        return None

    # Use storage directory for temp files (more reliable space than /tmp)
    temp_dir = _get_doc_temp_dir()

    try:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx",
             "--outdir", str(temp_dir), str(doc_path)],
            capture_output=True,
            timeout=timeout,
        )
        if result.returncode != 0:
            logger.warning(f"LibreOffice conversion failed for {doc_path}: {result.stderr}")
            return None

        # Find the converted file
        docx_name = doc_path.stem + ".docx"
        docx_path = temp_dir / docx_name
        if docx_path.exists():
            return docx_path

        logger.warning(f"Converted file not found: {docx_path}")
    except subprocess.TimeoutExpired:
        logger.warning(f"LibreOffice conversion timed out for {doc_path}")
    except Exception as e:
        logger.warning(f"Error converting {doc_path}: {e}")

    return None


def split_docx_into_heading_documents(
    docx_path: Path,
    ctx: "FileProcessingContext | None" = None
) -> List[LlamaIndexDocument]:
    """Split DOCX into documents by heading with progress updates.

    Args:
        docx_path: Path to DOCX file
        ctx: Optional processing context for progress updates and timeout
    """
    docs: List[LlamaIndexDocument] = []

    if ctx:
        ctx.set_phase("Opening DOCX")

    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.warning(f"Failed to open DOCX {docx_path}: {e}")
        return docs

    # Extract file dates from filesystem
    stat = docx_path.stat()
    creation_date = datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d")
    last_modified_date = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d")

    # Try to extract dates from DOCX core properties (more accurate than filesystem)
    try:
        core_props = doc.core_properties
        if core_props.created:
            creation_date = core_props.created.strftime("%Y-%m-%d")
        if core_props.modified:
            last_modified_date = core_props.modified.strftime("%Y-%m-%d")
    except Exception:
        pass  # Fall back to filesystem dates

    # First pass: Extract all headings with positions for hierarchy metadata
    if ctx:
        ctx.set_phase("Extracting headings")

    all_headings = []
    char_position = 0
    for para in doc.paragraphs:
        # Periodically check for timeout during long operations
        if ctx and len(all_headings) % 100 == 0:
            ctx.set_phase(f"Extracting headings ({len(all_headings)} found)")

        style_name = getattr(para.style, "name", "") or ""
        is_heading = (
            style_name.startswith("Heading")
            or style_name.startswith("heading")
            or "Heading" in style_name
        )

        if is_heading and para.text.strip():
            heading_level = _parse_heading_level(style_name)
            all_headings.append({
                "text": para.text.strip(),
                "position": char_position,
                "level": heading_level
            })

        char_position += len(para.text) + 1  # +1 for newline

    # Store headings separately to avoid metadata size issues during chunking
    if ctx:
        ctx.set_phase("Storing heading metadata")
    get_heading_store().set_headings(str(docx_path), all_headings)

    # Second pass: Split by heading (existing logic)
    if ctx:
        ctx.set_phase("Splitting into sections")
    current_heading: str | None = None
    current_level: int | None = None
    current_body: list[str] = []

    def flush_current():
        if not current_heading:
            return
        text = "\n".join(line for line in current_body if line is not None).strip()
        if not text:
            return

        # Build hierarchical heading_path by finding parent headings based on level
        heading_path = []
        if all_headings:
            # Find the index of the current heading in all_headings
            current_idx = None
            for idx, h in enumerate(all_headings):
                if h["text"] == current_heading and h["level"] == current_level:
                    current_idx = idx
                    break

            if current_idx is not None:
                # Build path by including all parent headings (those with lower level numbers)
                # Walk backwards from current heading and include headings with level < current_level
                path_headings = [all_headings[current_idx]]  # Start with current
                for idx in range(current_idx - 1, -1, -1):
                    h = all_headings[idx]
                    if h["level"] < path_headings[0]["level"]:
                        path_headings.insert(0, h)
                heading_path = [h["text"] for h in path_headings]

        metadata = {
            "file_path": str(docx_path),
            "file_name": docx_path.name,
            "source": str(docx_path),
            "heading": current_heading,
            "heading_level": current_level,
            "creation_date": creation_date,
            "last_modified_date": last_modified_date,
            "heading_path": heading_path,  # Pre-computed hierarchical path
        }
        docs.append(LlamaIndexDocument(
            text=text,
            metadata=metadata,
            excluded_embed_metadata_keys=EXCLUDED_EMBED_METADATA_KEYS,
            excluded_llm_metadata_keys=EXCLUDED_LLM_METADATA_KEYS,
        ))

    for para in doc.paragraphs:
        style_name = getattr(para.style, "name", "") or ""
        is_heading = (
            style_name.startswith("Heading")
            or style_name.startswith("heading")
            or "Heading" in style_name
        )

        if is_heading and para.text.strip():
            flush_current()
            current_heading = para.text.strip()
            current_level = _parse_heading_level(style_name)
            current_body = []
        else:
            if current_heading is not None:
                current_body.append(para.text)

    flush_current()

    if not docs:
        try:
            full_text = "\n".join(p.text for p in doc.paragraphs).strip()
        except Exception:
            full_text = ""

        if full_text:
            metadata = {
                "file_path": str(docx_path),
                "file_name": docx_path.name,
                "source": str(docx_path),
                "heading": None,
                "heading_level": None,
                "creation_date": creation_date,
                "last_modified_date": last_modified_date,
            }
            docs.append(LlamaIndexDocument(
                text=full_text,
                metadata=metadata,
                excluded_embed_metadata_keys=EXCLUDED_EMBED_METADATA_KEYS,
                excluded_llm_metadata_keys=EXCLUDED_LLM_METADATA_KEYS,
            ))

    logger.info(
        f"Split DOCX {docx_path} into {len(docs)} heading-based document(s)"
    )
    return docs


def tokenize_filename(filename: str) -> List[str]:
    """
    Tokenize a filename for BM25 indexing.

    Splits on delimiters (underscore, hyphen, dot, space) and camelCase.

    Examples:
        'cpp_styleguide.md' -> ['cpp', 'styleguide', 'md']
        'API-Reference-v2.pdf' -> ['api', 'reference', 'v2', 'pdf']
        'CamelCaseDoc.docx' -> ['camel', 'case', 'doc', 'docx']
    """
    import re

    name_parts = filename.rsplit('.', 1)
    base_name = name_parts[0]
    extension = name_parts[1] if len(name_parts) > 1 else ""

    # Split on explicit delimiters
    parts = re.split(r'[_\-\.\s]+', base_name)

    # Split camelCase within each part
    tokens = []
    for part in parts:
        camel_split = re.sub(r'([a-z])([A-Z])', r'\1 \2', part).split()
        tokens.extend(t.lower() for t in camel_split if t)

    # Add extension as a token
    if extension:
        tokens.append(extension.lower())

    return tokens


def build_bm25_index(index, storage_dir: Path) -> None:
    """
    Build a BM25 index over file names from the docstore.

    This enables keyword matching for queries like 'cpp styleguide' to find
    files named 'cpp_styleguide.md'.
    """
    from llama_index.retrievers.bm25 import BM25Retriever
    from llama_index.core.schema import TextNode

    logger.info("Building BM25 index for file name matching...")

    # Create filename nodes - one per unique file
    filename_nodes = []
    seen_files: Set[str] = set()

    for doc_id, node in index.docstore.docs.items():
        metadata = node.metadata or {}
        file_name = metadata.get("file_name", "")
        file_path = metadata.get("file_path", "")

        if not file_name or file_path in seen_files:
            continue
        seen_files.add(file_path)

        tokens = tokenize_filename(file_name)
        filename_nodes.append(TextNode(
            text=" ".join(tokens),
            metadata={"file_name": file_name, "file_path": file_path},
            id_=f"bm25_{file_path}"
        ))

    if not filename_nodes:
        logger.warning("No documents found for BM25 indexing")
        return

    logger.info(f"Creating BM25 index with {len(filename_nodes)} file name entries")

    bm25_retriever = BM25Retriever.from_defaults(
        nodes=filename_nodes,
        similarity_top_k=10,
    )

    bm25_dir = storage_dir / "bm25_index"
    bm25_dir.mkdir(parents=True, exist_ok=True)
    bm25_retriever.persist(str(bm25_dir))

    logger.info(f"BM25 index persisted to {bm25_dir}")


def batched(iterable, n):
    """Batch data into lists of length n. The last batch may be shorter.

    Compatible with Python <3.12 (which added itertools.batched).
    """
    it = iter(iterable)
    while True:
        batch = list(itertools.islice(it, n))
        if not batch:
            return
        yield batch


def configure_offline_mode(offline: bool, cache_dir: Path) -> None:
    """Configure environment variables for offline mode."""
    if offline:
        os.environ["HF_HUB_OFFLINE"] = "1"
        os.environ["TRANSFORMERS_OFFLINE"] = "1"
        os.environ["HF_DATASETS_OFFLINE"] = "1"
        cache_dir_abs = cache_dir.resolve()
        os.environ["HF_HOME"] = str(cache_dir_abs)
        os.environ["HF_HUB_CACHE"] = str(cache_dir_abs)
        os.environ["HF_DATASETS_CACHE"] = str(cache_dir_abs)
        logger.info("Offline mode enabled.")
    else:
        # Clear offline mode environment variables to allow downloads
        for var in ["HF_HUB_OFFLINE", "TRANSFORMERS_OFFLINE", "HF_DATASETS_OFFLINE"]:
            os.environ.pop(var, None)

    # Update huggingface_hub's cached constant (it caches at import time)
    try:
        from huggingface_hub import constants
        constants.HF_HUB_OFFLINE = offline
    except ImportError:
        pass


def calculate_optimal_batch_size(
    num_files: int,
    avg_chunks_per_file: int = 10,
    max_memory_mb: int = 2048,
    embedding_dims: int = 384,
) -> int:
    """Calculate optimal batch size to stay within memory budget.

    Args:
        num_files: Total number of files to process
        avg_chunks_per_file: Estimated average chunks per file
        max_memory_mb: Maximum memory budget in MB
        embedding_dims: Embedding dimensions (384 for BAAI/bge-small-en-v1.5)

    Returns:
        Optimal batch size (number of files to process together)
    """
    # Per-chunk memory estimate: embedding vector + source text + metadata overhead.
    # Embedding: dims × 4 bytes (float32) ≈ 1.5 KB for 384 dims.
    # Source text: ~2 KB avg per chunk (512 tokens ≈ ~2 KB).
    # LlamaIndex node overhead (metadata, doc_id, relationships): ~2 KB.
    # Total: ~5.5 KB per chunk. Use 6 KB as conservative estimate.
    bytes_per_chunk = max(embedding_dims * 4, 6 * 1024)

    total_chunks = num_files * avg_chunks_per_file
    total_memory_mb = (total_chunks * bytes_per_chunk) / (1024 * 1024)

    if total_memory_mb <= max_memory_mb:
        return num_files
    else:
        max_chunks = int((max_memory_mb * 1024 * 1024) / bytes_per_chunk)
        batch_size = max(1, int(max_chunks / avg_chunks_per_file))
        return min(batch_size, num_files)


def load_files_parallel(
    files: List[FileInfo],
    data_source: DataSource,
    ui: IndexingUI,
    max_workers: int = 4,
    timeout_enabled: bool = True,
    per_file_timeout: float = 300,
    heartbeat_interval: float = 2.0,
    abort_ctl: "GracefulAbort | None" = None,
) -> Dict[str, Tuple[FileInfo, List[Any]]]:
    """Load files in parallel using ThreadPoolExecutor.

    Args:
        files: List of FileInfo objects to load
        data_source: DataSource instance for loading files
        ui: IndexingUI instance for progress tracking
        max_workers: Maximum number of parallel workers
        timeout_enabled: Whether to enable per-file timeout
        per_file_timeout: Timeout in seconds per file
        heartbeat_interval: Heartbeat interval in seconds
        abort_ctl: Optional GracefulAbort instance for Ctrl-C handling

    Returns:
        Dict mapping file path to (FileInfo, List[LlamaIndexDocument])
    """
    file_docs = {}

    def load_single_file(file_info: FileInfo):
        """Load a single file (executed in thread pool)."""
        try:
            with FileProcessingContext(
                file_info.path,
                ui,
                timeout_seconds=per_file_timeout if timeout_enabled else None,
                heartbeat_interval=heartbeat_interval
            ) as ctx:
                ctx.set_phase("Loading file")
                docs = data_source.load_file(file_info, ctx)
                return file_info, docs
        except FileProcessingTimeoutError as e:
            logger.error(f"Timeout loading {file_info.path}: {e}")
            return file_info, []
        except Exception as e:
            logger.error(f"Error loading {file_info.path}: {e}")
            return file_info, []

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all file loading tasks
        future_to_file = {
            executor.submit(load_single_file, file_info): file_info
            for file_info in files
        }

        # Collect results as they complete
        for future in as_completed(future_to_file):
            if abort_ctl and abort_ctl.abort_requested:
                for f in future_to_file:
                    f.cancel()
                break

            file_info = future_to_file[future]
            try:
                result_file_info, docs = future.result(timeout=per_file_timeout + 10)
                if docs:  # Only add if we got documents
                    file_docs[result_file_info.path] = (result_file_info, docs)
            except Exception as e:
                logger.error(f"Error retrieving result for {file_info.path}: {e}")
            finally:
                ui.progress_update()

    return file_docs


class GracefulAbort:
    """Two-stage Ctrl-C handler for the indexing pipeline.

    First Ctrl-C:  sets abort flag and restores default SIGINT so a second
                   Ctrl-C raises KeyboardInterrupt immediately.
    """

    def __init__(self, ui: "IndexingUI"):
        self._abort = False
        self._ui = ui
        self._original_handler = None

    @property
    def abort_requested(self) -> bool:
        return self._abort

    def install(self) -> None:
        """Register the custom SIGINT handler. Must be called from the main thread."""
        self._original_handler = signal.getsignal(signal.SIGINT)
        signal.signal(signal.SIGINT, self._handle_sigint)

    def uninstall(self) -> None:
        """Restore the original SIGINT handler."""
        if self._original_handler is not None:
            signal.signal(signal.SIGINT, self._original_handler)
            self._original_handler = None

    def _handle_sigint(self, signum, frame):
        self._abort = True
        signal.signal(signal.SIGINT, signal.SIG_DFL)
        self._ui.print(
            f"\n{self._ui.YELLOW}Ctrl-C received. "
            f"Finishing current batch and saving state...{self._ui.RESET}"
        )
        self._ui.print(
            f"{self._ui.DIM}(Press Ctrl-C again to force quit immediately)"
            f"{self._ui.RESET}"
        )


def build_index(
    download_only: bool = False,
    config_path: Path | None = None,
    model_cache_dir: Path | None = None,
    verbose: bool = False,
) -> None:
    """Build and persist the vector index incrementally."""
    global _config, STORAGE_DIR, STATE_DB_PATH, RETRIEVAL_MODEL_CACHE_DIR, BM25_INDEX_DIR, HEADING_STORE_PATH
    global RETRIEVAL_EMBED_MODEL_NAME, RETRIEVAL_RERANK_MODEL_NAME

    if config_path:
        cfg = load_config(config_path)
        _config = cfg
        STORAGE_DIR = Path(cfg["storage"]["storage_dir"])
        STATE_DB_PATH = STORAGE_DIR / "ingestion_state.db"
        RETRIEVAL_MODEL_CACHE_DIR = Path(cfg["storage"]["model_cache_dir"])
        BM25_INDEX_DIR = STORAGE_DIR / "bm25_index"
        HEADING_STORE_PATH = STORAGE_DIR / "heading_store.json"
        RETRIEVAL_EMBED_MODEL_NAME = cfg["retrieval"]["embed_model_name"]
        RETRIEVAL_RERANK_MODEL_NAME = cfg["retrieval"]["rerank_model_name"]

    # Override model cache dir if specified via CLI
    if model_cache_dir:
        RETRIEVAL_MODEL_CACHE_DIR = model_cache_dir

    # Read offline setting from config; force online when downloading models
    offline = False if download_only else _config["retrieval"].get("offline", False)
    cache_dir = RETRIEVAL_MODEL_CACHE_DIR
    configure_offline_mode(offline, cache_dir)

    # Load configuration
    index_config = load_index_config()

    with IndexingUI(verbose=verbose) as ui:
        abort_ctl = GracefulAbort(ui)
        abort_ctl.install()

        # Model caching
        ui.step_start("Checking embedding model cache")
        ensure_embedding_model_cached(cache_dir, offline=offline)
        ui.step_done()

        ui.step_start("Checking rerank model cache")
        try:
            ensure_rerank_model_cached(cache_dir, offline=offline)
            ui.step_done()
        except FileNotFoundError:
            if download_only or offline:
                raise
            ui.step_done("skipped")

        if download_only:
            ui.success("Models downloaded successfully.")
            return

        # Initialize State and Multi-Directory Data Source
        ingestion_state = IngestionState(STATE_DB_PATH)
        data_source = MultiDirectoryDataSource(index_config)

        # Directory summary
        summary = data_source.get_summary()
        dir_msg = f"Directories: {len(summary['available'])} available"
        if summary['unavailable']:
            dir_msg += f", {len(summary['unavailable'])} unavailable"
        ui.print(dir_msg)

        if not data_source.sources:
            ui.error("Error: No available directories to index. Check your config.yaml.")
            return

        # Initialize Embedding Model
        ui.step_start("Initializing embedding model")
        embed_model = _create_fastembed_embedding(RETRIEVAL_MODEL_CACHE_DIR, offline=offline)
        Settings.embed_model = embed_model
        ui.step_done()

        if abort_ctl.abort_requested:
            abort_ctl.uninstall()
            ui.print(f"{ui.YELLOW}Interrupted during initialization. No changes made.{ui.RESET}")
            return

        # Configure Text Splitter using config values
        text_splitter = SentenceSplitter(
            chunk_size=index_config.chunk_size,
            chunk_overlap=index_config.chunk_overlap,
            separator=" ",
        )
        Settings.text_splitter = text_splitter

        # Load existing index or create new
        if (STORAGE_DIR / "docstore.json").exists():
            ui.step_start("Loading existing index")
            storage_context = StorageContext.from_defaults(persist_dir=str(STORAGE_DIR))
            index = load_index_from_storage(storage_context, embed_model=embed_model)
            ui.step_done()
        else:
            ui.step_start("Creating new index")
            storage_context = StorageContext.from_defaults()
            index = VectorStoreIndex([], storage_context=storage_context, embed_model=embed_model)
            ui.step_done()

        # Change Detection
        tracked_files = ingestion_state.get_all_files()
        found_files: Set[str] = set()
        files_to_process: List[FileInfo] = []
        new_count = 0
        modified_count = 0

        ui.step_start("Scanning for changes")
        scan_count = 0
        for file_info in data_source.iter_files(tracked_files=tracked_files):
            if abort_ctl.abort_requested:
                break
            scan_count += 1
            if scan_count % 100 == 0:
                ui.step_update(f"Scanning for changes ({scan_count:,} files)")
            found_files.add(file_info.path)
            existing_state = tracked_files.get(file_info.path)

            if existing_state:
                if existing_state["hash"] != file_info.hash:
                    modified_count += 1
                    files_to_process.append(file_info)
            else:
                new_count += 1
                files_to_process.append(file_info)

        if abort_ctl.abort_requested:
            ui.step_done("interrupted")
            abort_ctl.uninstall()
            ui.print(f"{ui.YELLOW}Interrupted during scan. No changes made.{ui.RESET}")
            return

        # Identify Deleted Files
        deleted_files = set(tracked_files.keys()) - found_files

        # Build summary suffix
        parts = []
        if new_count:
            parts.append(f"{new_count} new")
        if modified_count:
            parts.append(f"{modified_count} modified")
        if deleted_files:
            parts.append(f"{len(deleted_files)} deleted")
        ui.step_done(", ".join(parts) if parts else "no changes")

        # Process deletions
        for deleted_path in deleted_files:
            doc_ids = tracked_files[deleted_path]["doc_ids"]
            for doc_id in doc_ids:
                try:
                    index.delete_ref_doc(doc_id, delete_from_docstore=True)
                except Exception as e:
                    logger.warning(f"Failed to delete doc {doc_id} from index: {e}")
            get_heading_store().remove_headings(deleted_path)
            ingestion_state.remove_file_state(deleted_path)

        if not files_to_process and not deleted_files:
            ui.success("Index is up to date.")
            return

        # Process New/Modified Files
        if files_to_process:
            # Get configuration
            timeout_enabled = cfgload.get("indexing.timeout.enabled", True)
            per_file_timeout = cfgload.get("indexing.timeout.per_file_seconds", 300)
            heartbeat_interval = cfgload.get("indexing.timeout.heartbeat_interval_seconds", 2)

            # Checkpointing configuration (controls how often we persist to disk)
            checkpoint_interval_files = cfgload.get("indexing.checkpoint_interval_files", 500)
            checkpoint_interval_seconds = cfgload.get("indexing.checkpoint_interval_seconds", 300)

            # Batch sizing (controls how many files are loaded/embedded together)
            configured_batch_size = cfgload.get("indexing.batch_size", 200)
            enable_adaptive_batching = cfgload.get("indexing.enable_adaptive_batching", True)
            max_memory_mb = cfgload.get("indexing.max_memory_mb", 2048)

            if enable_adaptive_batching:
                optimal_batch_size = calculate_optimal_batch_size(
                    num_files=len(files_to_process),
                    max_memory_mb=max_memory_mb
                )
                batch_size = min(configured_batch_size, optimal_batch_size)
            else:
                batch_size = configured_batch_size

            # Delete old versions of modified files
            ui.step_start("Removing old versions of modified files")
            for file_info in files_to_process:
                existing_state = tracked_files.get(file_info.path)
                if existing_state:
                    for doc_id in existing_state["doc_ids"]:
                        try:
                            index.delete_ref_doc(doc_id, delete_from_docstore=True)
                        except KeyError:
                            pass
            ui.step_done()

            # Process files in batches with checkpointing
            total_files = len(files_to_process)
            files_since_checkpoint = 0
            last_checkpoint_time = time.time()

            ui.progress_start(total_files, desc="Processing files", unit="file")

            # Get parallel loading configuration
            max_workers = cfgload.get("indexing.parallel_workers", 4)
            enable_parallel = cfgload.get("indexing.enable_parallel_loading", True)

            batch_num = 0
            total_batches = (total_files + batch_size - 1) // batch_size

            for batch in batched(files_to_process, batch_size):
                if abort_ctl.abort_requested:
                    break

                batch_num += 1

                # Load files in this batch
                accumulated_docs = []
                doc_to_file_mapping = {}
                file_doc_ids = {}

                workers = max_workers if enable_parallel else 1
                file_docs = load_files_parallel(
                    batch,
                    data_source,
                    ui,
                    max_workers=workers,
                    timeout_enabled=timeout_enabled,
                    per_file_timeout=per_file_timeout,
                    heartbeat_interval=heartbeat_interval,
                    abort_ctl=abort_ctl,
                )

                for file_path, (file_info, docs) in file_docs.items():
                    file_doc_ids[file_path] = []
                    for doc in docs:
                        accumulated_docs.append(doc)
                        doc_to_file_mapping[doc.doc_id] = file_info
                        file_doc_ids[file_path].append(doc.doc_id)

                # Batch embedding
                if accumulated_docs:
                    ui.progress_substep_start(f"Converting {len(accumulated_docs)} documents to nodes")
                    nodes = text_splitter.get_nodes_from_documents(accumulated_docs)
                    ui.progress_substep_done()

                    batch_label = f" (batch {batch_num}/{total_batches})" if total_batches > 1 else ""
                    ui.progress_substep_start(f"Embedding {len(nodes)} nodes{batch_label}")

                    # Pre-compute embeddings in bulk via FastEmbed ONNX directly,
                    # bypassing LlamaIndex's per-batch instrumentation overhead.
                    # embed_nodes() skips nodes where node.embedding is already set.
                    texts = [
                        node.get_content(metadata_mode=MetadataMode.EMBED)
                        for node in nodes
                    ]
                    embeddings = embed_model._get_text_embeddings(texts)
                    for node, emb in zip(nodes, embeddings):
                        node.embedding = emb

                    index.insert_nodes(nodes)
                    ui.progress_substep_done()

                # Batch update state
                state_updates = []
                for file_path, doc_ids in file_doc_ids.items():
                    if doc_ids:
                        file_info = doc_to_file_mapping[doc_ids[0]]
                        state_updates.append((file_info, doc_ids))

                if state_updates:
                    ingestion_state.update_file_states_batch(state_updates)

                # Checkpoint only when interval thresholds are met
                files_since_checkpoint += len(batch)
                elapsed = time.time() - last_checkpoint_time
                should_checkpoint = (
                    files_since_checkpoint >= checkpoint_interval_files
                    or elapsed >= checkpoint_interval_seconds
                )
                if should_checkpoint:
                    index.storage_context.persist(persist_dir=str(STORAGE_DIR))
                    with get_heading_store() as store:
                        pass  # Context manager will auto-flush
                    files_since_checkpoint = 0
                    last_checkpoint_time = time.time()

                if abort_ctl.abort_requested:
                    break

            # Handle graceful abort after batch loop
            if abort_ctl.abort_requested:
                ui.progress_done()
                ui.step_start("Saving checkpoint before exit")
                STORAGE_DIR.mkdir(parents=True, exist_ok=True)
                index.storage_context.persist(persist_dir=str(STORAGE_DIR))
                get_heading_store().flush()
                ui.step_done()
                processed = ui._progress_current
                abort_ctl.uninstall()
                ui.print(
                    f"{ui.YELLOW}Indexing interrupted. "
                    f"{processed} of {total_files} files processed. "
                    f"Run again to continue.{ui.RESET}"
                )
                return

        # Final checkpoint
        STORAGE_DIR.mkdir(parents=True, exist_ok=True)
        ui.step_start("Persisting final index")
        index.storage_context.persist(persist_dir=str(STORAGE_DIR))
        get_heading_store().flush()
        ui.step_done()

        ui.step_start("Building BM25 index")
        build_bm25_index(index, STORAGE_DIR)
        ui.step_done()

        abort_ctl.uninstall()
        ui.success(f"\nIndexing complete. {len(files_to_process)} files processed.")


if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    )
    parser = argparse.ArgumentParser(description="Build the document index")
    parser.add_argument(
        "--download-models",
        action="store_true",
        help="Download the retrieval models and exit",
    )
    parser.add_argument(
        "--config",
        type=str,
        help="Path to config.yaml (overrides auto-discovery)",
    )
    parser.add_argument(
        "--model-cache-dir",
        type=str,
        help="Directory to download/cache models (overrides config)",
    )
    args = parser.parse_args()

    try:
        build_index(
            download_only=args.download_models,
            config_path=Path(args.config) if args.config else None,
            model_cache_dir=Path(args.model_cache_dir) if args.model_cache_dir else None,
        )
    except Exception as e:
        logger.error(f"Indexing failed: {e}", exc_info=True)
        raise
