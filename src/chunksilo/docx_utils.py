#!/usr/bin/env python3
# SPDX-License-Identifier: Apache-2.0
"""DOCX/DOC document processing utilities.

Handles parsing DOCX files into heading-based documents, heading level
extraction, and .doc-to-.docx conversion via LibreOffice.
"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any

from docx import Document
from llama_index.core import Document as LlamaIndexDocument

from . import cfgload

if TYPE_CHECKING:
    from .ui import FileProcessingContext

logger = logging.getLogger(__name__)


def _parse_heading_level(style_name: str | None) -> int:
    """Best-effort extraction of a numeric heading level from a DOCX style name."""
    if not style_name:
        return 1
    try:
        if "Heading" in style_name:
            level_str = style_name.replace("Heading", "").strip()
            if level_str:
                return int(level_str)
    except (ValueError, AttributeError):
        pass
    return 1


def _get_doc_temp_dir() -> Path:
    """Get the temporary directory for .doc conversion, creating it if needed."""
    storage_dir = Path(cfgload.get("storage.storage_dir", "./storage"))
    temp_dir = storage_dir / "doc_temp"
    temp_dir.mkdir(parents=True, exist_ok=True)
    return temp_dir


def _convert_doc_to_docx(doc_path: Path, timeout: float = 60) -> Path | None:
    """Convert a .doc file to .docx using LibreOffice.

    Args:
        doc_path: Path to .doc file
        timeout: Timeout in seconds for conversion process

    Returns:
        Path to temporary .docx file, or None if conversion fails.
        Caller is responsible for cleaning up the temp file.
    """
    import shutil
    import subprocess

    # Find LibreOffice executable
    soffice_paths = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
        "/usr/bin/soffice",  # Linux
        "/usr/bin/libreoffice",  # Linux alternative
        "soffice",  # Windows (in PATH)
    ]

    soffice = None
    for path in soffice_paths:
        if shutil.which(path):
            soffice = path
            break

    if not soffice:
        logger.warning(f"LibreOffice not found. Cannot convert {doc_path}")
        return None

    # Use storage directory for temp files (more reliable space than /tmp)
    temp_dir = _get_doc_temp_dir()

    try:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx",
             "--outdir", str(temp_dir), str(doc_path)],
            capture_output=True,
            timeout=timeout,
        )
        if result.returncode != 0:
            logger.warning(f"LibreOffice conversion failed for {doc_path}: {result.stderr}")
            return None

        # Find the converted file
        docx_name = doc_path.stem + ".docx"
        docx_path = temp_dir / docx_name
        if docx_path.exists():
            return docx_path

        logger.warning(f"Converted file not found: {docx_path}")
    except subprocess.TimeoutExpired:
        logger.warning(f"LibreOffice conversion timed out for {doc_path}")
    except Exception as e:
        logger.warning(f"Error converting {doc_path}: {e}")

    return None


def split_docx_into_heading_documents(
    docx_path: Path,
    ctx: FileProcessingContext | None = None,
    *,
    heading_store: Any = None,
    excluded_embed_metadata_keys: list[str] | None = None,
    excluded_llm_metadata_keys: list[str] | None = None,
) -> list[LlamaIndexDocument]:
    """Split DOCX into documents by heading with progress updates.

    Args:
        docx_path: Path to DOCX file
        ctx: Optional processing context for progress updates and timeout
        heading_store: HeadingStore instance for persisting heading metadata
        excluded_embed_metadata_keys: Keys to exclude from embedding text
        excluded_llm_metadata_keys: Keys to exclude from LLM context
    """
    _excluded_embed = excluded_embed_metadata_keys or []
    _excluded_llm = excluded_llm_metadata_keys or []

    docs: list[LlamaIndexDocument] = []

    if ctx:
        ctx.set_phase("Opening DOCX")

    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.warning(f"Failed to open DOCX {docx_path}: {e}")
        return docs

    # Extract file dates from filesystem
    stat = docx_path.stat()
    creation_date = datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d")
    last_modified_date = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d")

    # Try to extract dates from DOCX core properties (more accurate than filesystem)
    try:
        core_props = doc.core_properties
        if core_props.created:
            creation_date = core_props.created.strftime("%Y-%m-%d")
        if core_props.modified:
            last_modified_date = core_props.modified.strftime("%Y-%m-%d")
    except Exception:
        pass  # Fall back to filesystem dates

    # First pass: Extract all headings with positions for hierarchy metadata
    if ctx:
        ctx.set_phase("Extracting headings")

    all_headings = []
    char_position = 0
    for para in doc.paragraphs:
        # Periodically check for timeout during long operations
        if ctx and len(all_headings) % 100 == 0:
            ctx.set_phase(f"Extracting headings ({len(all_headings)} found)")

        style_name = getattr(para.style, "name", "") or ""
        is_heading = (
            style_name.startswith("Heading")
            or style_name.startswith("heading")
            or "Heading" in style_name
        )

        if is_heading and para.text.strip():
            heading_level = _parse_heading_level(style_name)
            all_headings.append({
                "text": para.text.strip(),
                "position": char_position,
                "level": heading_level
            })

        char_position += len(para.text) + 1  # +1 for newline

    # Store headings separately to avoid metadata size issues during chunking
    if heading_store is not None:
        if ctx:
            ctx.set_phase("Storing heading metadata")
        heading_store.set_headings(str(docx_path), all_headings)

    # Second pass: Split by heading (existing logic)
    if ctx:
        ctx.set_phase("Splitting into sections")
    current_heading: str | None = None
    current_level: int | None = None
    current_body: list[str] = []

    def flush_current():
        if not current_heading:
            return
        text = "\n".join(line for line in current_body if line is not None).strip()
        if not text:
            return

        # Build hierarchical heading_path by finding parent headings based on level
        heading_path = []
        if all_headings:
            # Find the index of the current heading in all_headings
            current_idx = None
            for idx, h in enumerate(all_headings):
                if h["text"] == current_heading and h["level"] == current_level:
                    current_idx = idx
                    break

            if current_idx is not None:
                # Build path by including all parent headings (those with lower level numbers)
                # Walk backwards from current heading and include headings with level < current_level
                path_headings = [all_headings[current_idx]]  # Start with current
                for idx in range(current_idx - 1, -1, -1):
                    h = all_headings[idx]
                    if h["level"] < path_headings[0]["level"]:
                        path_headings.insert(0, h)
                heading_path = [h["text"] for h in path_headings]

        metadata = {
            "file_path": str(docx_path),
            "file_name": docx_path.name,
            "source": str(docx_path),
            "heading": current_heading,
            "heading_level": current_level,
            "creation_date": creation_date,
            "last_modified_date": last_modified_date,
            "heading_path": heading_path,  # Pre-computed hierarchical path
        }
        docs.append(LlamaIndexDocument(
            text=text,
            metadata=metadata,
            excluded_embed_metadata_keys=_excluded_embed,
            excluded_llm_metadata_keys=_excluded_llm,
        ))

    for para in doc.paragraphs:
        style_name = getattr(para.style, "name", "") or ""
        is_heading = (
            style_name.startswith("Heading")
            or style_name.startswith("heading")
            or "Heading" in style_name
        )

        if is_heading and para.text.strip():
            flush_current()
            current_heading = para.text.strip()
            current_level = _parse_heading_level(style_name)
            current_body = []
        else:
            if current_heading is not None:
                current_body.append(para.text)

    flush_current()

    if not docs:
        try:
            full_text = "\n".join(p.text for p in doc.paragraphs).strip()
        except Exception:
            full_text = ""

        if full_text:
            metadata = {
                "file_path": str(docx_path),
                "file_name": docx_path.name,
                "source": str(docx_path),
                "heading": None,
                "heading_level": None,
                "creation_date": creation_date,
                "last_modified_date": last_modified_date,
            }
            docs.append(LlamaIndexDocument(
                text=full_text,
                metadata=metadata,
                excluded_embed_metadata_keys=_excluded_embed,
                excluded_llm_metadata_keys=_excluded_llm,
            ))

    logger.info(
        f"Split DOCX {docx_path} into {len(docs)} heading-based document(s)"
    )
    return docs
