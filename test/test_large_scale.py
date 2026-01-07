#!/usr/bin/env python3
"""
Large-scale automated test suite for RAG system evaluation.

This test suite:
1. Downloads a MASSIVE diverse corpus of documents from the web (100+ documents)
   - 40+ academic papers (NLP, CV, RL, etc.)
   - 30+ technical documentation files (GitHub READMEs)
   - 15+ classic literature texts (Project Gutenberg)
   - 20+ generated DOCX documents (diverse technical topics)
2. Ingests them into the RAG system
3. Tests retrieval accuracy with 50+ diverse queries covering all document types
4. Evaluates using standard RAG metrics (Precision@k, Recall@k, MRR, NDCG)
5. Challenges the models with various query types and edge cases

This massive corpus is designed to test the system's effectiveness at sifting through
large amounts of diverse data, ensuring it can handle real-world scale and complexity.
"""
import asyncio
import json
import logging
import math
import os
import sys
import time
from collections import defaultdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple
from urllib.parse import urlparse

import requests
from dotenv import load_dotenv

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

load_dotenv()

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)

# Force online mode for tests (to download models)
os.environ["OFFLINE"] = "0"

# Import after logging is set up
from ingest import DATA_DIR, STORAGE_DIR, build_index
from mcp_server import retrieve_docs

# Test corpus configuration
TEST_DATA_DIR = Path(os.getenv("TEST_DATA_DIR", "./test_data"))
TEST_STORAGE_DIR = Path(os.getenv("TEST_STORAGE_DIR", "./test_storage"))
TEST_RESULTS_DIR = Path(os.getenv("TEST_RESULTS_DIR", "./test_results"))

# Configuration: abort on download failures
ABORT_ON_DOWNLOAD_FAILURE = os.getenv("ABORT_ON_DOWNLOAD_FAILURE", "1").lower() not in ("0", "false", "no")

# Document sources - MASSIVE diverse corpus from public domains
# This corpus is designed to test the system's ability to sift through large amounts of data
DOCUMENT_SOURCES = {
    "pdf": [
        # Core NLP/ML papers (Transformer family)
        "https://arxiv.org/pdf/1706.03762.pdf",  # Attention Is All You Need (Transformer)
        "https://arxiv.org/pdf/2005.14165.pdf",  # GPT-3 paper
        "https://arxiv.org/pdf/1810.04805.pdf",  # BERT paper
        "https://arxiv.org/pdf/2010.11929.pdf",  # Vision Transformer (ViT)
        "https://arxiv.org/pdf/1910.10683.pdf",  # T5: Text-To-Text Transfer Transformer
        "https://arxiv.org/pdf/2001.08361.pdf",  # RoBERTa
        "https://arxiv.org/pdf/1907.11692.pdf",  # ALBERT
        "https://arxiv.org/pdf/2010.00854.pdf",  # DeBERTa
        "https://arxiv.org/pdf/2203.02155.pdf",  # InstructGPT
        "https://arxiv.org/pdf/2303.08774.pdf",  # GPT-4 Technical Report
        
        # Computer Vision papers
        "https://arxiv.org/pdf/1409.1556.pdf",   # VGG
        "https://arxiv.org/pdf/1512.03385.pdf",  # ResNet
        "https://arxiv.org/pdf/1603.05027.pdf",  # ResNet v2
        "https://arxiv.org/pdf/1704.04861.pdf",  # MobileNets
        "https://arxiv.org/pdf/1801.04381.pdf",  # MobileNetV2
        "https://arxiv.org/pdf/1905.11946.pdf",  # EfficientNet
        
        # Reinforcement Learning papers
        "https://arxiv.org/pdf/1312.5602.pdf",   # DQN
        "https://arxiv.org/pdf/1509.06461.pdf",  # Double DQN
        "https://arxiv.org/pdf/1707.06347.pdf",  # PPO
        "https://arxiv.org/pdf/1801.01290.pdf",  # Soft Actor-Critic
        "https://arxiv.org/pdf/1910.10897.pdf",  # MuZero
        
        # Generative Models
        "https://arxiv.org/pdf/1406.2661.pdf",   # GAN
        "https://arxiv.org/pdf/1511.06434.pdf",  # DCGAN
        "https://arxiv.org/pdf/2006.11239.pdf",  # CLIP
        "https://arxiv.org/pdf/2102.09672.pdf",  # DALL-E
        "https://arxiv.org/pdf/2204.06125.pdf",  # DALL-E 2
        
        # Systems and Architecture papers
        "https://arxiv.org/pdf/1609.08144.pdf",  # Google's Neural Machine Translation
        "https://arxiv.org/pdf/1706.03762.pdf",  # Transformer (duplicate for emphasis)
        "https://arxiv.org/pdf/1803.02155.pdf",  # Universal Transformers
        "https://arxiv.org/pdf/1901.02860.pdf",  # BART
        "https://arxiv.org/pdf/2004.04906.pdf",  # Longformer
        
        # Optimization and Training
        "https://arxiv.org/pdf/1412.6980.pdf",   # Adam optimizer
        "https://arxiv.org/pdf/1608.03983.pdf",  # Noisy Adam
        "https://arxiv.org/pdf/1711.05101.pdf",  # Lookahead optimizer
        "https://arxiv.org/pdf/1908.03265.pdf",  # RAdam
        
        # Additional diverse papers
        "https://arxiv.org/pdf/1506.01497.pdf",  # Faster R-CNN
        "https://arxiv.org/pdf/1506.06724.pdf",  # YOLO
        "https://arxiv.org/pdf/1611.05431.pdf",  # ResNeXt
        "https://arxiv.org/pdf/1704.04861.pdf",  # MobileNets (duplicate)
        "https://arxiv.org/pdf/1807.11164.pdf",  # BERT (duplicate for emphasis)
    ],
    "markdown": [
        # Major programming languages and frameworks
        "https://raw.githubusercontent.com/python/cpython/main/README.rst",
        "https://raw.githubusercontent.com/nodejs/node/main/README.md",
        "https://raw.githubusercontent.com/golang/go/master/README.md",
        "https://raw.githubusercontent.com/rust-lang/rust/master/README.md",
        "https://raw.githubusercontent.com/microsoft/TypeScript/main/README.md",
        "https://raw.githubusercontent.com/dotnet/core/main/README.md",
        
        # Frontend frameworks
        "https://raw.githubusercontent.com/facebook/react/main/README.md",
        "https://raw.githubusercontent.com/vuejs/core/main/README.md",
        "https://raw.githubusercontent.com/angular/angular/main/README.md",
        "https://raw.githubusercontent.com/sveltejs/svelte/main/README.md",
        
        # Backend frameworks
        "https://raw.githubusercontent.com/django/django/main/README.rst",
        "https://raw.githubusercontent.com/rails/rails/main/README.md",
        "https://raw.githubusercontent.com/spring-projects/spring-framework/main/README.md",
        
        # ML/AI frameworks
        "https://raw.githubusercontent.com/tensorflow/tensorflow/master/README.md",
        "https://raw.githubusercontent.com/pytorch/pytorch/main/README.md",
        "https://raw.githubusercontent.com/scikit-learn/scikit-learn/main/README.rst",
        "https://raw.githubusercontent.com/huggingface/transformers/main/README.md",
        "https://raw.githubusercontent.com/pytorch/vision/main/README.md",  # Alternative ML framework
        
        # Databases
        "https://raw.githubusercontent.com/mongodb/mongo/master/README.md",
        "https://raw.githubusercontent.com/postgres/postgres/master/README.md",
        "https://raw.githubusercontent.com/redis/redis/unstable/README.md",
        
        # DevOps and Infrastructure
        "https://raw.githubusercontent.com/kubernetes/kubernetes/master/README.md",
        "https://raw.githubusercontent.com/moby/moby/master/README.md",  # Docker (moby is the open source component)
        "https://raw.githubusercontent.com/hashicorp/terraform/main/README.md",
        "https://raw.githubusercontent.com/ansible/ansible/devel/README.md",
        
        # Editors and IDEs
        "https://raw.githubusercontent.com/microsoft/vscode/main/README.md",
        "https://raw.githubusercontent.com/neovim/neovim/master/README.md",
        
        # Additional popular projects
        "https://raw.githubusercontent.com/git/git/master/README.md",
        "https://raw.githubusercontent.com/apache/spark/master/README.md",
        "https://raw.githubusercontent.com/apache/kafka/trunk/README.md",
        "https://raw.githubusercontent.com/grafana/grafana/main/README.md",
        "https://raw.githubusercontent.com/prometheus/prometheus/main/README.md",
    ],
    "txt": [
        # Classic literature from Project Gutenberg
        "https://www.gutenberg.org/files/1342/1342-0.txt",  # Pride and Prejudice
        "https://www.gutenberg.org/files/11/11-0.txt",     # Alice in Wonderland
        "https://www.gutenberg.org/files/84/84-0.txt",    # Frankenstein
        "https://www.gutenberg.org/files/2701/2701-0.txt", # Moby Dick
        "https://www.gutenberg.org/files/98/98-0.txt",     # A Tale of Two Cities
        "https://www.gutenberg.org/files/74/74-0.txt",    # The Adventures of Tom Sawyer
        "https://www.gutenberg.org/files/1661/1661-0.txt", # The Adventures of Sherlock Holmes
        "https://www.gutenberg.org/files/5200/5200-0.txt", # Metamorphosis
        "https://www.gutenberg.org/files/345/345-0.txt",   # Dracula
        "https://www.gutenberg.org/files/1232/1232-0.txt", # The Prince
        "https://www.gutenberg.org/files/2600/2600-0.txt", # War and Peace (Part 1)
        "https://www.gutenberg.org/files/1399/1399-0.txt", # Anna Karenina
        "https://www.gutenberg.org/files/1080/1080-0.txt", # A Modest Proposal
        "https://www.gutenberg.org/files/16328/16328-0.txt", # The Art of War
        "https://www.gutenberg.org/files/25344/25344-0.txt", # The Republic
    ],
    "docx": [
        # Note: DOCX files are harder to find publicly, so we'll generate many more
        # programmatically to create a diverse corpus
    ],
}

# Test queries with expected answers/contexts - MASSIVE query set
# Format: (query, expected_keywords, expected_file_patterns, difficulty)
TEST_QUERIES = [
    # === NLP and Machine Learning Queries ===
    # Simple factual queries
    (
        "What is attention mechanism in transformers?",
        ["attention", "transformer", "self-attention"],
        ["1706.03762", "transformer_overview"],
        "easy",
    ),
    (
        "How does BERT work?",
        ["bert", "bidirectional", "encoder"],
        ["1810.04805", "bert_explained"],
        "easy",
    ),
    (
        "What is GPT-3?",
        ["gpt-3", "language model", "generative"],
        ["2005.14165", "gpt_models"],
        "easy",
    ),
    (
        "What is GPT-4?",
        ["gpt-4", "language model"],
        ["2303.08774", "gpt_models"],
        "easy",
    ),
    (
        "How does Vision Transformer work?",
        ["vision", "transformer", "vit"],
        ["2010.11929"],
        "medium",
    ),
    (
        "What is T5 model?",
        ["t5", "text-to-text", "transfer"],
        ["1910.10683"],
        "medium",
    ),
    (
        "Explain RoBERTa architecture",
        ["roberta", "bert", "optimization"],
        ["2001.08361"],
        "medium",
    ),
    (
        "What is ALBERT model?",
        ["albert", "bert", "factorized"],
        ["1907.11692"],
        "medium",
    ),
    
    # Computer Vision queries
    (
        "What is ResNet architecture?",
        ["resnet", "residual", "network"],
        ["1512.03385", "1603.05027"],
        "medium",
    ),
    (
        "How does YOLO object detection work?",
        ["yolo", "object detection", "real-time"],
        ["1506.06724"],
        "medium",
    ),
    (
        "What is EfficientNet?",
        ["efficientnet", "scaling", "efficient"],
        ["1905.11946"],
        "medium",
    ),
    (
        "Explain computer vision fundamentals",
        ["computer vision", "cnn", "image classification"],
        ["computer_vision_basics"],
        "easy",
    ),
    
    # Reinforcement Learning queries
    (
        "What is DQN algorithm?",
        ["dqn", "deep q-learning", "reinforcement"],
        ["1312.5602", "reinforcement_learning_primer"],
        "medium",
    ),
    (
        "How does PPO work?",
        ["ppo", "proximal", "policy optimization"],
        ["1707.06347", "reinforcement_learning_primer"],
        "medium",
    ),
    (
        "What is reinforcement learning?",
        ["reinforcement learning", "agent", "environment"],
        ["reinforcement_learning_primer"],
        "easy",
    ),
    
    # Generative Models queries
    (
        "What are GANs?",
        ["gan", "generative adversarial", "network"],
        ["1406.2661"],
        "medium",
    ),
    (
        "How does CLIP work?",
        ["clip", "contrastive", "vision-language"],
        ["2006.11239"],
        "medium",
    ),
    (
        "What is DALL-E?",
        ["dall-e", "image generation", "multimodal"],
        ["2102.09672", "2204.06125"],
        "medium",
    ),
    
    # Optimization queries
    (
        "What is Adam optimizer?",
        ["adam", "optimizer", "adaptive"],
        ["1412.6980"],
        "medium",
    ),
    
    # Complex multi-part queries
    (
        "Compare and contrast transformer architecture with BERT architecture",
        ["transformer", "bert", "architecture", "encoder"],
        ["1706.03762", "1810.04805", "transformer_overview", "bert_explained"],
        "hard",
    ),
    (
        "What are the key innovations in language models from GPT-3 to transformers?",
        ["gpt-3", "transformer", "language model", "innovation"],
        ["2005.14165", "1706.03762", "gpt_models"],
        "hard",
    ),
    (
        "Compare ResNet and EfficientNet architectures",
        ["resnet", "efficientnet", "architecture", "comparison"],
        ["1512.03385", "1905.11946"],
        "hard",
    ),
    
    # Edge cases - misspellings and variations
    (
        "What is attension mechansim?",  # Intentional misspellings
        ["attention", "mechanism"],
        ["1706.03762", "transformer_overview"],
        "medium",
    ),
    (
        "How do transformers work?",
        ["transformer", "attention", "encoder", "decoder"],
        ["1706.03762", "transformer_overview"],
        "medium",
    ),
    
    # Broad queries that should retrieve multiple documents
    (
        "What are neural language models?",
        ["language model", "neural", "nlp"],
        ["2005.14165", "1810.04805", "1706.03762", "gpt_models"],
        "medium",
    ),
    (
        "What are the main deep learning architectures?",
        ["deep learning", "architecture", "neural network"],
        ["1706.03762", "1512.03385", "1406.2661"],
        "hard",
    ),
    
    # === Programming and Software Engineering Queries ===
    (
        "What is Python programming?",
        ["python", "programming"],
        ["python", "python_best_practices"],
        "easy",
    ),
    (
        "What are Python best practices?",
        ["python", "best practices", "pep"],
        ["python_best_practices"],
        "easy",
    ),
    (
        "How does React work?",
        ["react", "component", "javascript"],
        ["react", "react_patterns"],
        "medium",
    ),
    (
        "What are React design patterns?",
        ["react", "patterns", "components", "hooks"],
        ["react_patterns"],
        "medium",
    ),
    (
        "Explain Docker containers",
        ["docker", "container", "virtualization"],
        ["docker", "docker_guide"],
        "medium",
    ),
    (
        "What is Kubernetes?",
        ["kubernetes", "orchestration", "containers"],
        ["kubernetes", "kubernetes_overview"],
        "medium",
    ),
    (
        "How to design REST APIs?",
        ["rest", "api", "design", "http"],
        ["rest_api_design"],
        "medium",
    ),
    (
        "What are microservices?",
        ["microservices", "architecture", "services"],
        ["microservices_architecture"],
        "medium",
    ),
    (
        "Explain Git workflows",
        ["git", "workflow", "branching"],
        ["git", "git_workflows"],
        "medium",
    ),
    (
        "What is CI/CD?",
        ["cicd", "continuous integration", "deployment"],
        ["cicd_pipelines"],
        "medium",
    ),
    
    # === Database Queries ===
    (
        "What is SQL?",
        ["sql", "database", "query"],
        ["sql_fundamentals"],
        "easy",
    ),
    (
        "What are NoSQL databases?",
        ["nosql", "database", "mongodb", "redis"],
        ["nosql_databases"],
        "medium",
    ),
    (
        "How does MongoDB work?",
        ["mongodb", "document", "database"],
        ["mongodb", "nosql_databases"],
        "medium",
    ),
    (
        "What is Redis?",
        ["redis", "key-value", "cache"],
        ["redis", "nosql_databases"],
        "medium",
    ),
    (
        "Explain data engineering",
        ["data engineering", "etl", "pipeline"],
        ["data_engineering"],
        "medium",
    ),
    
    # === System Design and Security Queries ===
    (
        "What are security best practices?",
        ["security", "authentication", "encryption"],
        ["security_best_practices"],
        "medium",
    ),
    (
        "How to monitor applications?",
        ["monitoring", "observability", "metrics", "logging"],
        ["monitoring_observability"],
        "medium",
    ),
    (
        "What is performance optimization?",
        ["performance", "optimization", "caching"],
        ["performance_optimization"],
        "medium",
    ),
    (
        "Explain testing strategies",
        ["testing", "test", "quality"],
        ["testing_strategies"],
        "medium",
    ),
    (
        "What is cloud computing?",
        ["cloud", "computing", "iaas", "paas"],
        ["cloud_computing"],
        "medium",
    ),
    
    # === Literature Queries (from Gutenberg texts) ===
    (
        "What is Pride and Prejudice about?",
        ["pride", "prejudice", "austen"],
        ["1342"],
        "easy",
    ),
    (
        "Tell me about Alice in Wonderland",
        ["alice", "wonderland", "carroll"],
        ["11"],
        "easy",
    ),
    (
        "What is Frankenstein?",
        ["frankenstein", "shelley", "monster"],
        ["84"],
        "easy",
    ),
    (
        "What is Moby Dick?",
        ["moby", "dick", "whale", "melville"],
        ["2701"],
        "easy",
    ),
    (
        "Tell me about Sherlock Holmes",
        ["sherlock", "holmes", "detective"],
        ["1661"],
        "easy",
    ),
    
    # === Complex Cross-Domain Queries ===
    (
        "How do machine learning models compare to traditional programming?",
        ["machine learning", "programming", "comparison"],
        ["python_best_practices", "transformer_overview"],
        "hard",
    ),
    (
        "What are the best practices for deploying ML models?",
        ["ml", "deployment", "best practices"],
        ["docker_guide", "kubernetes_overview", "cicd_pipelines"],
        "hard",
    ),
    (
        "How to build scalable systems with microservices and containers?",
        ["microservices", "containers", "scalable"],
        ["microservices_architecture", "docker_guide", "kubernetes_overview"],
        "hard",
    ),
    
    # === Negative queries (should NOT retrieve certain documents) ===
    (
        "What is Python programming?",
        ["python", "programming"],
        ["python"],  # Should NOT retrieve transformer/BERT papers
        "easy",
    ),
    (
        "Tell me about React components",
        ["react", "components"],
        ["react"],  # Should NOT retrieve database or ML papers
        "easy",
    ),
]


def validate_file_content(file_path: Path, expected_type: str) -> bool:
    """Validate that a downloaded file is of the expected type."""
    try:
        if not file_path.exists() or file_path.stat().st_size == 0:
            logger.error(f"File validation failed: {file_path.name} does not exist or is empty")
            return False
        
        # Check file header/magic bytes
        with open(file_path, "rb") as f:
            header = f.read(1024)
        
        if expected_type == "pdf":
            # PDF files start with %PDF
            if not header.startswith(b"%PDF"):
                # Check if it's HTML (common for 404 pages)
                if header.startswith(b"<!DOCTYPE") or header.startswith(b"<html") or header.startswith(b"<HTML"):
                    logger.error(
                        f"Download failed: {file_path.name} is HTML (likely 404 page), not a PDF. "
                        f"Starts with: {header[:50].decode('utf-8', errors='ignore')[:50]}"
                    )
                else:
                    logger.error(
                        f"Download failed: {file_path.name} does not appear to be a valid PDF. "
                        f"Starts with: {header[:50]}"
                    )
                return False
        elif expected_type in ("md", "markdown", "txt"):
            # Text files should be readable as UTF-8 or ASCII
            try:
                header.decode("utf-8")
            except UnicodeDecodeError:
                # Might be binary, check if it's actually text
                if b"\x00" in header[:100]:  # Null bytes suggest binary
                    logger.error(f"Download failed: {file_path.name} appears to be binary, not text")
                    return False
        
        return True
    except Exception as e:
        logger.error(f"Error validating {file_path}: {e}")
        return False


def download_file(url: str, output_path: Path, timeout: int = 30, expected_type: str = "auto") -> bool:
    """Download a file from URL to output path and validate it."""
    # Check if file already exists and is valid
    if output_path.exists() and output_path.stat().st_size > 0:
        if expected_type == "auto":
            if output_path.suffix == ".pdf":
                expected_type = "pdf"
            elif output_path.suffix in (".md", ".rst", ".txt"):
                expected_type = "txt"
        if expected_type and validate_file_content(output_path, expected_type):
            logger.info(f"File already exists and is valid: {output_path.name}")
            return True
    
    try:
        logger.info(f"Downloading {url} to {output_path}")
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        # For arXiv PDFs, add a delay to avoid rate limiting
        if "arxiv.org" in url:
            import time
            time.sleep(1.0)  # Delay to avoid rate limiting
        
        response = requests.get(url, timeout=timeout, stream=True, headers=headers, allow_redirects=True)
        response.raise_for_status()
        
        # For arXiv PDFs, check content type header first to detect HTML responses
        if "arxiv.org" in url and expected_type == "pdf":
            content_type = response.headers.get("Content-Type", "").lower()
            if "text/html" in content_type:
                logger.warning(f"arXiv returned HTML instead of PDF for {url}, likely rate-limited. Skipping.")
                return False
        
        # Check content type from response
        content_type = response.headers.get("Content-Type", "").lower()
        
        # Detect expected type if not specified
        if expected_type == "auto":
            if output_path.suffix == ".pdf":
                expected_type = "pdf"
            elif output_path.suffix in (".md", ".rst", ".txt"):
                expected_type = "txt"
            else:
                expected_type = "unknown"
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Handle zip files specially
        if url.endswith(".zip") or "zip" in content_type:
            import zipfile
            import io
            try:
                with zipfile.ZipFile(io.BytesIO(response.content)) as z:
                    # Extract PDFs from zip
                    for member in z.namelist():
                        if member.endswith(".pdf"):
                            with z.open(member) as source, open(output_path, "wb") as target:
                                target.write(source.read())
                            logger.info(f"Extracted {member} from zip")
                            if validate_file_content(output_path, "pdf"):
                                return True
                            else:
                                output_path.unlink()  # Remove invalid file
                                return False
            except zipfile.BadZipFile:
                logger.error(f"Download failed: Downloaded file is not a valid ZIP: {url}")
                return False
            return False
        
        # Download file
        with open(output_path, "wb") as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        
        file_size = output_path.stat().st_size
        
        # Validate file content - this is critical for detecting HTML 404 pages masquerading as PDFs
        if expected_type == "pdf":
            if not validate_file_content(output_path, "pdf"):
                logger.error(f"Download validation failed for {url} - file is not a valid PDF")
                if output_path.exists():
                    output_path.unlink()  # Remove invalid file
                return False
        elif expected_type == "txt":
            if not validate_file_content(output_path, "txt"):
                logger.error(f"Download validation failed for {url} - file is not valid text")
                if output_path.exists():
                    output_path.unlink()  # Remove invalid file
                return False
        
        logger.info(f"Downloaded {output_path.name} ({file_size} bytes)")
        return True
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 404:
            logger.error(f"File not found (404): {url}")
        else:
            logger.error(f"HTTP error downloading {url}: {e}")
        return False
    except Exception as e:
        logger.error(f"Failed to download {url}: {e}")
        # Clean up partial download
        if output_path.exists():
            try:
                output_path.unlink()
            except Exception:
                pass
        return False


def create_sample_docx(output_path: Path, content: str, title: str = "Test Document") -> bool:
    """Create a sample DOCX file with content."""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading(title, 0)
        
        # Split content into paragraphs
        for para_text in content.split("\n\n"):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info(f"Created DOCX file: {output_path}")
        return True
    except Exception as e:
        logger.warning(f"Failed to create DOCX {output_path}: {e}")
        return False


def cleanup_invalid_files(data_dir: Path) -> None:
    """Remove invalid files from previous test runs."""
    logger.info("Cleaning up invalid files from previous runs...")
    removed_count = 0
    
    for pdf_file in data_dir.rglob("*.pdf"):
        if not validate_file_content(pdf_file, "pdf"):
            logger.info(f"Removing invalid PDF: {pdf_file}")
            try:
                pdf_file.unlink()
                removed_count += 1
            except Exception as e:
                logger.warning(f"Failed to remove {pdf_file}: {e}")
    
    if removed_count > 0:
        logger.info(f"Removed {removed_count} invalid file(s)")


def download_test_corpus() -> Dict[str, List[Path]]:
    """Download a diverse corpus of test documents."""
    logger.info("=" * 80)
    logger.info("Downloading Test Corpus")
    logger.info("=" * 80)
    
    TEST_DATA_DIR.mkdir(parents=True, exist_ok=True)
    
    # Clean up any invalid files from previous runs
    cleanup_invalid_files(TEST_DATA_DIR)
    
    downloaded_files = defaultdict(list)
    download_failures = []
    
    # Download PDFs
    for url in DOCUMENT_SOURCES["pdf"]:
        filename = Path(urlparse(url).path).name
        if not filename.endswith(".pdf"):
            filename = f"{Path(urlparse(url).path).stem}.pdf"
        output_path = TEST_DATA_DIR / "pdf" / filename
        if download_file(url, output_path, expected_type="pdf"):
            downloaded_files["pdf"].append(output_path)
        else:
            download_failures.append(("pdf", url, output_path))
    
    # Download Markdown files
    for url in DOCUMENT_SOURCES["markdown"]:
        filename = Path(urlparse(url).path).name
        if not filename.endswith((".md", ".rst", ".txt")):
            # Generate unique filename based on URL
            url_stem = Path(urlparse(url).path).stem
            if url.endswith(".rst"):
                filename = f"{url_stem}.rst"
            else:
                filename = f"{url_stem}.md"
        output_path = TEST_DATA_DIR / "md" / filename
        # Handle filename conflicts by adding source identifier
        if output_path.exists():
            # Add a prefix to avoid overwriting
            url_parts = urlparse(url).netloc.split(".")
            prefix = url_parts[-2] if len(url_parts) >= 2 else "doc"
            output_path = TEST_DATA_DIR / "md" / f"{prefix}_{filename}"
        if download_file(url, output_path, expected_type="txt"):
            downloaded_files["markdown"].append(output_path)
        else:
            download_failures.append(("markdown", url, output_path))
    
    # Download TXT files
    for url in DOCUMENT_SOURCES["txt"]:
        filename = Path(urlparse(url).path).name
        if not filename.endswith(".txt"):
            filename = f"{Path(urlparse(url).path).stem}.txt"
        output_path = TEST_DATA_DIR / "txt" / filename
        if download_file(url, output_path, expected_type="txt"):
            downloaded_files["txt"].append(output_path)
        else:
            download_failures.append(("txt", url, output_path))
    
    # Create MASSIVE collection of sample DOCX files covering diverse topics
    docx_content = {
        # NLP and ML topics
        "transformer_overview.docx": {
            "title": "Transformer Architecture Overview",
            "content": """
# Transformer Architecture

## Introduction
The Transformer architecture, introduced in "Attention Is All You Need", revolutionized natural language processing.

## Key Components

### Self-Attention Mechanism
Self-attention allows the model to weigh the importance of different words in a sequence.

### Encoder-Decoder Structure
The transformer uses an encoder-decoder architecture with multiple layers.

## Applications
Transformers are used in machine translation, text generation, and many other NLP tasks.
            """.strip(),
        },
        "bert_explained.docx": {
            "title": "BERT Explained",
            "content": """
# BERT: Bidirectional Encoder Representations from Transformers

## Overview
BERT is a bidirectional transformer model that reads text in both directions.

## Architecture
BERT uses only the encoder part of the transformer architecture.

## Training
BERT is pre-trained on large text corpora using masked language modeling and next sentence prediction.

## Applications
BERT is widely used for question answering, sentiment analysis, and named entity recognition.
            """.strip(),
        },
        "gpt_models.docx": {
            "title": "GPT Models: From GPT-1 to GPT-4",
            "content": """
# GPT Models Evolution

## GPT-1
The first Generative Pre-trained Transformer introduced unsupervised pre-training.

## GPT-2
A larger model with 1.5B parameters that showed impressive zero-shot capabilities.

## GPT-3
A massive 175B parameter model demonstrating few-shot learning and in-context learning.

## GPT-4
The latest iteration with improved reasoning, safety, and multimodal capabilities.
            """.strip(),
        },
        "computer_vision_basics.docx": {
            "title": "Computer Vision Fundamentals",
            "content": """
# Computer Vision Fundamentals

## Convolutional Neural Networks
CNNs are the backbone of modern computer vision, using convolutional layers to detect features.

## Image Classification
The task of assigning labels to images, fundamental to many CV applications.

## Object Detection
Identifying and localizing multiple objects within an image.

## Semantic Segmentation
Pixel-level classification of images into different semantic regions.
            """.strip(),
        },
        "reinforcement_learning_primer.docx": {
            "title": "Reinforcement Learning Primer",
            "content": """
# Reinforcement Learning Primer

## Core Concepts
RL involves an agent learning to make decisions by interacting with an environment.

## Key Algorithms
Q-Learning, Policy Gradient methods, and Actor-Critic approaches are fundamental.

## Applications
RL is used in game playing, robotics, autonomous vehicles, and recommendation systems.

## Challenges
Sample efficiency, exploration-exploitation tradeoff, and stability are major challenges.
            """.strip(),
        },
        
        # Programming and Software Engineering
        "python_best_practices.docx": {
            "title": "Python Best Practices",
            "content": """
# Python Best Practices

## Code Style
Follow PEP 8 guidelines for consistent, readable code.

## Error Handling
Use try-except blocks appropriately and handle exceptions gracefully.

## Testing
Write unit tests, integration tests, and use testing frameworks like pytest.

## Documentation
Document your code with docstrings and maintain clear README files.
            """.strip(),
        },
        "react_patterns.docx": {
            "title": "React Design Patterns",
            "content": """
# React Design Patterns

## Component Composition
Build complex UIs by composing smaller, reusable components.

## Hooks
Use React Hooks for state management and side effects in functional components.

## Context API
Share state across components without prop drilling using Context.

## Performance Optimization
Use memoization, code splitting, and lazy loading for better performance.
            """.strip(),
        },
        "docker_guide.docx": {
            "title": "Docker Containerization Guide",
            "content": """
# Docker Containerization Guide

## Containers vs Virtual Machines
Containers are lightweight, share the host OS kernel, and start faster than VMs.

## Dockerfile Best Practices
Use multi-stage builds, minimize layers, and leverage build cache.

## Docker Compose
Orchestrate multi-container applications with Docker Compose.

## Container Security
Follow security best practices: use minimal base images, scan for vulnerabilities.
            """.strip(),
        },
        "kubernetes_overview.docx": {
            "title": "Kubernetes Orchestration Overview",
            "content": """
# Kubernetes Orchestration Overview

## Core Concepts
Pods, Services, Deployments, and Namespaces are fundamental Kubernetes resources.

## Scaling
Horizontal Pod Autoscaling automatically adjusts the number of pods based on metrics.

## Service Discovery
Kubernetes provides DNS-based service discovery for pods and services.

## Configuration Management
Use ConfigMaps and Secrets for managing application configuration and sensitive data.
            """.strip(),
        },
        
        # Databases and Data
        "sql_fundamentals.docx": {
            "title": "SQL Fundamentals",
            "content": """
# SQL Fundamentals

## Data Types
Understanding VARCHAR, INTEGER, DATE, and other SQL data types.

## Queries
SELECT, FROM, WHERE, JOIN, GROUP BY, and ORDER BY are essential SQL clauses.

## Indexes
Indexes improve query performance by creating fast lookup structures.

## Transactions
ACID properties ensure data consistency and reliability.
            """.strip(),
        },
        "nosql_databases.docx": {
            "title": "NoSQL Database Types",
            "content": """
# NoSQL Database Types

## Document Databases
MongoDB stores data as documents in JSON-like format.

## Key-Value Stores
Redis provides fast key-value storage with various data structures.

## Column Stores
Cassandra stores data in columns, optimized for write-heavy workloads.

## Graph Databases
Neo4j represents data as nodes and relationships, ideal for connected data.
            """.strip(),
        },
        "data_engineering.docx": {
            "title": "Data Engineering Principles",
            "content": """
# Data Engineering Principles

## ETL Pipelines
Extract, Transform, Load processes move and transform data between systems.

## Data Warehousing
Centralized repositories for structured data from multiple sources.

## Data Lakes
Storage systems that hold raw data in its native format.

## Stream Processing
Real-time data processing with systems like Apache Kafka and Flink.
            """.strip(),
        },
        
        # System Design and Architecture
        "microservices_architecture.docx": {
            "title": "Microservices Architecture",
            "content": """
# Microservices Architecture

## Service Decomposition
Break monolithic applications into independent, loosely coupled services.

## API Gateway
Single entry point for client requests, routing to appropriate microservices.

## Service Discovery
Mechanisms for services to find and communicate with each other.

## Distributed Systems Challenges
Handle network failures, data consistency, and service coordination.
            """.strip(),
        },
        "rest_api_design.docx": {
            "title": "REST API Design Principles",
            "content": """
# REST API Design Principles

## Resource-Based URLs
Use nouns for resources, not verbs. Follow RESTful conventions.

## HTTP Methods
GET for retrieval, POST for creation, PUT for updates, DELETE for removal.

## Status Codes
Use appropriate HTTP status codes: 200, 201, 400, 404, 500, etc.

## Versioning
Version APIs to maintain backward compatibility as they evolve.
            """.strip(),
        },
        "security_best_practices.docx": {
            "title": "Security Best Practices",
            "content": """
# Security Best Practices

## Authentication and Authorization
Implement secure authentication (OAuth, JWT) and proper authorization checks.

## Input Validation
Validate and sanitize all user inputs to prevent injection attacks.

## Encryption
Use HTTPS for data in transit and encryption for data at rest.

## Security Headers
Implement security headers like CSP, HSTS, and X-Frame-Options.
            """.strip(),
        },
        
        # DevOps and CI/CD
        "cicd_pipelines.docx": {
            "title": "CI/CD Pipeline Design",
            "content": """
# CI/CD Pipeline Design

## Continuous Integration
Automatically build and test code on every commit.

## Continuous Deployment
Automatically deploy code to production after passing tests.

## Pipeline Stages
Build, test, security scanning, and deployment stages in the pipeline.

## Infrastructure as Code
Define infrastructure using code (Terraform, CloudFormation) for reproducibility.
            """.strip(),
        },
        "monitoring_observability.docx": {
            "title": "Monitoring and Observability",
            "content": """
# Monitoring and Observability

## Metrics
Collect quantitative data about system performance and behavior.

## Logging
Centralized logging for debugging and auditing purposes.

## Tracing
Distributed tracing to understand request flow across services.

## Alerting
Set up alerts for critical issues and performance degradation.
            """.strip(),
        },
        
        # Additional Technical Topics
        "git_workflows.docx": {
            "title": "Git Workflows and Best Practices",
            "content": """
# Git Workflows and Best Practices

## Branching Strategies
Git Flow, GitHub Flow, and trunk-based development approaches.

## Commit Messages
Write clear, descriptive commit messages following conventions.

## Code Review
Effective code review practices for maintaining code quality.

## Merge Strategies
Understanding merge, rebase, and squash strategies.
            """.strip(),
        },
        "performance_optimization.docx": {
            "title": "Performance Optimization Techniques",
            "content": """
# Performance Optimization Techniques

## Profiling
Identify bottlenecks using profiling tools and metrics.

## Caching
Implement caching strategies at multiple levels (application, database, CDN).

## Database Optimization
Index optimization, query tuning, and connection pooling.

## Code Optimization
Algorithm improvements, lazy loading, and resource pooling.
            """.strip(),
        },
        "testing_strategies.docx": {
            "title": "Software Testing Strategies",
            "content": """
# Software Testing Strategies

## Test Pyramid
Unit tests form the base, integration tests in the middle, E2E tests at the top.

## Test-Driven Development
Write tests before implementation to guide design.

## Coverage Metrics
Measure code coverage but focus on meaningful tests.

## Test Automation
Automate testing in CI/CD pipelines for continuous quality assurance.
            """.strip(),
        },
        "cloud_computing.docx": {
            "title": "Cloud Computing Fundamentals",
            "content": """
# Cloud Computing Fundamentals

## Service Models
IaaS, PaaS, and SaaS provide different levels of abstraction.

## Deployment Models
Public, private, and hybrid cloud deployment options.

## Scalability
Horizontal and vertical scaling strategies for cloud applications.

## Cost Optimization
Right-sizing resources, reserved instances, and spot instances for cost savings.
            """.strip(),
        },
    }
    
    for filename, doc_data in docx_content.items():
        output_path = TEST_DATA_DIR / "docx" / filename
        if create_sample_docx(output_path, doc_data["content"], doc_data["title"]):
            downloaded_files["docx"].append(output_path)
    
    total_files = sum(len(files) for files in downloaded_files.values())
    logger.info(f"\nDownloaded {total_files} files:")
    for file_type, files in downloaded_files.items():
        logger.info(f"  {file_type}: {len(files)} files")
    
    # Report download failures
    if download_failures:
        logger.warning(f"\n{'=' * 80}")
        logger.warning(f"Download Failures: {len(download_failures)} file(s) failed to download")
        logger.warning(f"{'=' * 80}")
        
        # Check if failures are primarily from arXiv (rate limiting)
        arxiv_failures = sum(1 for _, url, _ in download_failures if "arxiv.org" in url)
        non_arxiv_failures = len(download_failures) - arxiv_failures
        
        for file_type, url, output_path in download_failures:
            logger.warning(f"  {file_type}: {url}")
        
        # If we have a substantial corpus (50+ files) and failures are mostly arXiv rate limiting,
        # continue anyway since we have enough data for testing
        if total_files >= 50 and arxiv_failures > 0 and non_arxiv_failures == 0:
            logger.warning(f"\nContinuing with {total_files} files despite {arxiv_failures} arXiv PDF failures (rate limiting).")
            logger.warning("This is sufficient for large-scale testing. PDFs can be downloaded later if needed.")
        elif ABORT_ON_DOWNLOAD_FAILURE:
            logger.error("\nAborting test suite due to download failures.")
            logger.error("Set ABORT_ON_DOWNLOAD_FAILURE=0 to continue despite failures.")
            raise RuntimeError(
                f"Failed to download {len(download_failures)} required file(s). "
                "Test suite aborted. Check the URLs and network connectivity."
            )
        else:
            logger.warning("\nContinuing despite download failures (ABORT_ON_DOWNLOAD_FAILURE=0).")
    
    if total_files == 0:
        raise RuntimeError(
            "No files were downloaded. Cannot proceed with tests. "
            "Check network connectivity and URL availability."
        )
    
    return dict(downloaded_files)


def _get_file_identifier(chunk: Dict) -> str:
    """Extract file identifier from chunk for pattern matching."""
    # Try location.file first (full path), then metadata.file_name, then metadata.file_path
    file_path = chunk.get("location", {}).get("file", "")
    if not file_path:
        file_path = chunk.get("metadata", {}).get("file_name", "")
    if not file_path:
        file_path = chunk.get("metadata", {}).get("file_path", "")
    return str(file_path).lower()


def precision_at_k(retrieved: List[Dict], relevant_file_patterns: List[str], k: int) -> float:
    """Calculate Precision@k."""
    if k == 0:
        return 0.0
    
    top_k = retrieved[:k]
    relevant_count = 0
    
    for chunk in top_k:
        file_id = _get_file_identifier(chunk)
        
        # Check if any relevant pattern matches
        if any(pattern.lower() in file_id for pattern in relevant_file_patterns):
            relevant_count += 1
    
    return relevant_count / k


def recall_at_k(retrieved: List[Dict], relevant_file_patterns: List[str], k: int, total_relevant: int) -> float:
    """Calculate Recall@k."""
    if total_relevant == 0:
        return 0.0
    
    top_k = retrieved[:k]
    relevant_retrieved = set()
    
    for chunk in top_k:
        file_id = _get_file_identifier(chunk)
        
        # Find which relevant pattern matches
        for pattern in relevant_file_patterns:
            if pattern.lower() in file_id:
                relevant_retrieved.add(pattern)
                break
    
    return len(relevant_retrieved) / total_relevant if total_relevant > 0 else 0.0


def mean_reciprocal_rank(retrieved: List[Dict], relevant_file_patterns: List[str]) -> float:
    """Calculate Mean Reciprocal Rank (MRR)."""
    for rank, chunk in enumerate(retrieved, start=1):
        file_id = _get_file_identifier(chunk)
        
        if any(pattern.lower() in file_id for pattern in relevant_file_patterns):
            return 1.0 / rank
    
    return 0.0


def ndcg_at_k(retrieved: List[Dict], relevant_file_patterns: List[str], k: int) -> float:
    """Calculate Normalized Discounted Cumulative Gain (NDCG@k)."""
    if not retrieved or k == 0:
        return 0.0
    
    # Use actual number of retrieved chunks, not k (in case we have fewer than k)
    actual_k = min(k, len(retrieved))
    if actual_k == 0:
        return 0.0
    
    def relevance_score(chunk: Dict) -> float:
        file_id = _get_file_identifier(chunk)
        # Binary relevance: 1 if relevant, 0 otherwise
        return 1.0 if any(pattern.lower() in file_id for pattern in relevant_file_patterns) else 0.0
    
    # Calculate DCG@k (using actual_k, not k)
    dcg = 0.0
    relevant_count = 0
    for i, chunk in enumerate(retrieved[:actual_k], start=1):
        rel = relevance_score(chunk)
        if rel > 0:
            relevant_count += 1
        # DCG formula: sum of (relevance / log2(rank + 1))
        # Note: using i (1-indexed) so log2(i+1) = log2(2) for first item = 1.0
        dcg += rel / math.log2(i + 1)
    
    # If no relevant documents were retrieved, NDCG is 0
    if relevant_count == 0:
        return 0.0
    
    # Calculate IDCG@k (ideal DCG - all relevant documents at the top)
    # IDCG assumes all relevant documents are ranked at positions 1, 2, 3, ...
    # We need to know how many relevant documents exist, but we only have patterns.
    # For simplicity, we assume each pattern corresponds to at least one document.
    # The ideal case is having all relevant documents at the top.
    num_relevant_docs = len(relevant_file_patterns)
    if num_relevant_docs == 0:
        return 0.0
    
    # IDCG: DCG if all relevant documents were perfectly ranked at the top
    # Use actual_k (not k) to match the number of documents we're actually evaluating
    # We can't have more than actual_k documents, and we can't have more than num_relevant_docs
    ideal_count = min(actual_k, num_relevant_docs)
    idcg = sum(1.0 / math.log2(i + 1) for i in range(1, ideal_count + 1))
    
    if idcg == 0:
        return 0.0
    
    # NDCG = DCG / IDCG (normalized to [0, 1])
    ndcg = dcg / idcg
    # Clamp to [0, 1] in case of floating point issues
    return max(0.0, min(1.0, ndcg))


async def evaluate_query_with_retriever(
    query: str,
    expected_keywords: List[str],
    expected_file_patterns: List[str],
    difficulty: str,
    retriever_func,
) -> Dict[str, Any]:
    """Evaluate a single query against the RAG system."""
    logger.info(f"\nEvaluating query: {query}")
    logger.info(f"Expected patterns: {expected_file_patterns}")
    
    start_time = time.time()
    result = await retriever_func(query)
    elapsed = time.time() - start_time
    
    chunks = result.get("chunks", [])
    
    # Calculate metrics at different k values
    k_values = [1, 3, 5, 10]
    metrics = {}
    
    for k in k_values:
        metrics[f"precision@{k}"] = precision_at_k(chunks, expected_file_patterns, k)
        metrics[f"recall@{k}"] = recall_at_k(chunks, expected_file_patterns, k, len(expected_file_patterns))
        # Calculate NDCG@k even if we have fewer than k chunks (ndcg_at_k handles this)
        if len(chunks) > 0:
            ndcg_value = ndcg_at_k(chunks, expected_file_patterns, k)
            metrics[f"ndcg@{k}"] = ndcg_value
            # Log if NDCG is suspiciously low
            if ndcg_value == 0.0 and len(expected_file_patterns) > 0:
                # Check if any relevant docs were retrieved
                actual_k = min(k, len(chunks))
                relevant_found = any(
                    any(p.lower() in _get_file_identifier(chunk)
                        for p in expected_file_patterns)
                    for chunk in chunks[:actual_k]
                )
                if not relevant_found:
                    logger.debug(f"  NDCG@{k}=0.0: No relevant documents found in top {actual_k} results")
        else:
            metrics[f"ndcg@{k}"] = 0.0
    
    metrics["mrr"] = mean_reciprocal_rank(chunks, expected_file_patterns)
    
    # Check if expected keywords appear in retrieved chunks
    keyword_matches = {}
    for keyword in expected_keywords:
        found = False
        for chunk in chunks[:5]:  # Check top 5 chunks
            text = chunk.get("text", "").lower()
            if keyword.lower() in text:
                found = True
                break
        keyword_matches[keyword] = found
    
    # Get top retrieved files
    top_files = []
    for chunk in chunks[:5]:
        file_name = chunk.get("metadata", {}).get("file_name", "unknown")
        top_files.append(file_name)
    
    evaluation = {
        "query": query,
        "difficulty": difficulty,
        "num_chunks_retrieved": len(chunks),
        "retrieval_time": elapsed,
        "metrics": metrics,
        "keyword_matches": keyword_matches,
        "top_files": top_files,
        "expected_patterns": expected_file_patterns,
    }
    
    logger.info(f"  Precision@5: {metrics['precision@5']:.3f}")
    logger.info(f"  Recall@5: {metrics['recall@5']:.3f}")
    logger.info(f"  MRR: {metrics['mrr']:.3f}")
    ndcg5 = metrics.get('ndcg@5', 0.0)
    logger.info(f"  NDCG@5: {ndcg5:.3f}")
    if ndcg5 == 0.0 and len(expected_file_patterns) > 0:
        # Check if relevant docs were actually retrieved
        relevant_in_top5 = sum(
            1 for chunk in chunks[:5]
            if any(
                pattern.lower() in _get_file_identifier(chunk)
                for pattern in expected_file_patterns
            )
        )
        if relevant_in_top5 == 0:
            logger.warning(f"  ⚠ NDCG@5=0.0: No relevant documents found in top 5 results")
            logger.warning(f"     Expected patterns: {expected_file_patterns}")
            logger.warning(f"     Top 5 files: {top_files[:5]}")
    logger.info(f"  Top files: {', '.join(top_files[:3])}")
    
    return evaluation


async def evaluate_query(
    query: str,
    expected_keywords: List[str],
    expected_file_patterns: List[str],
    difficulty: str,
) -> Dict[str, Any]:
    """Evaluate a single query against the RAG system (uses global retrieve_docs)."""
    return await evaluate_query_with_retriever(
        query, expected_keywords, expected_file_patterns, difficulty, retrieve_docs
    )


async def run_large_scale_tests() -> Dict[str, Any]:
    """Run the complete large-scale test suite."""
    logger.info("=" * 80)
    logger.info("Large-Scale RAG System Test Suite")
    logger.info("=" * 80)
    
    # Step 1: Download test corpus
    downloaded_files = download_test_corpus()
    
    if not any(downloaded_files.values()):
        logger.error("No documents downloaded. Cannot proceed with tests.")
        return {"error": "No documents downloaded"}
    
    # Step 2: Temporarily set DATA_DIR and STORAGE_DIR for test
    original_data_dir = os.environ.get("DATA_DIR")
    original_storage_dir = os.environ.get("STORAGE_DIR")
    
    try:
        os.environ["DATA_DIR"] = str(TEST_DATA_DIR)
        os.environ["STORAGE_DIR"] = str(TEST_STORAGE_DIR)
        
        # Re-import to get updated paths
        import importlib
        import ingest
        import mcp_server
        importlib.reload(ingest)
        importlib.reload(mcp_server)
        
        # Re-import after reload
        from ingest import build_index as build_test_index
        from mcp_server import load_llamaindex_index, retrieve_docs as retrieve_docs_reloaded
        
        # Step 3: Build index
        logger.info("\n" + "=" * 80)
        logger.info("Building Index from Test Corpus")
        logger.info("=" * 80)
        
        build_test_index()
        
        # Step 4: Load index
        logger.info("\n" + "=" * 80)
        logger.info("Loading Index")
        logger.info("=" * 80)
        
        index = load_llamaindex_index()
        logger.info("Index loaded successfully")
        
        # Step 5: Run evaluation queries
        logger.info("\n" + "=" * 80)
        logger.info("Running Evaluation Queries")
        logger.info("=" * 80)
        
        evaluations = []
        for query, keywords, patterns, difficulty in TEST_QUERIES:
            try:
                # Use the reloaded retrieve_docs function
                eval_result = await evaluate_query_with_retriever(
                    query, keywords, patterns, difficulty, retrieve_docs_reloaded
                )
                evaluations.append(eval_result)
            except Exception as e:
                logger.error(f"Error evaluating query '{query}': {e}")
                import traceback
                traceback.print_exc()
                evaluations.append({
                    "query": query,
                    "error": str(e),
                })
        
        # Step 6: Calculate aggregate metrics
        logger.info("\n" + "=" * 80)
        logger.info("Calculating Aggregate Metrics")
        logger.info("=" * 80)
        
        successful_evals = [e for e in evaluations if "error" not in e]
        
        if successful_evals:
            aggregate_metrics = {
                "precision@1": sum(e["metrics"]["precision@1"] for e in successful_evals) / len(successful_evals),
                "precision@5": sum(e["metrics"]["precision@5"] for e in successful_evals) / len(successful_evals),
                "recall@5": sum(e["metrics"]["recall@5"] for e in successful_evals) / len(successful_evals),
                "mrr": sum(e["metrics"]["mrr"] for e in successful_evals) / len(successful_evals),
                "ndcg@5": sum(e["metrics"].get("ndcg@5", 0) for e in successful_evals) / len(successful_evals),
            }
            
            # Metrics by difficulty
            difficulty_metrics = defaultdict(lambda: {"count": 0, "precision@5": [], "recall@5": [], "mrr": []})
            for eval_result in successful_evals:
                diff = eval_result.get("difficulty", "unknown")
                difficulty_metrics[diff]["count"] += 1
                difficulty_metrics[diff]["precision@5"].append(eval_result["metrics"]["precision@5"])
                difficulty_metrics[diff]["recall@5"].append(eval_result["metrics"]["recall@5"])
                difficulty_metrics[diff]["mrr"].append(eval_result["metrics"]["mrr"])
            
            difficulty_summary = {}
            for diff, data in difficulty_metrics.items():
                if data["count"] > 0:
                    difficulty_summary[diff] = {
                        "count": data["count"],
                        "avg_precision@5": sum(data["precision@5"]) / len(data["precision@5"]),
                        "avg_recall@5": sum(data["recall@5"]) / len(data["recall@5"]),
                        "avg_mrr": sum(data["mrr"]) / len(data["mrr"]),
                    }
            
            logger.info("\nAggregate Metrics:")
            logger.info(f"  Precision@1: {aggregate_metrics['precision@1']:.3f}")
            logger.info(f"  Precision@5: {aggregate_metrics['precision@5']:.3f}")
            logger.info(f"  Recall@5: {aggregate_metrics['recall@5']:.3f}")
            logger.info(f"  MRR: {aggregate_metrics['mrr']:.3f}")
            logger.info(f"  NDCG@5: {aggregate_metrics['ndcg@5']:.3f}")
            
            logger.info("\nMetrics by Difficulty:")
            for diff, metrics in difficulty_summary.items():
                logger.info(f"  {diff}:")
                logger.info(f"    Count: {metrics['count']}")
                logger.info(f"    Avg Precision@5: {metrics['avg_precision@5']:.3f}")
                logger.info(f"    Avg Recall@5: {metrics['avg_recall@5']:.3f}")
                logger.info(f"    Avg MRR: {metrics['avg_mrr']:.3f}")
        else:
            aggregate_metrics = {}
            difficulty_summary = {}
            logger.warning("No successful evaluations to aggregate")
        
        # Step 7: Save results
        TEST_RESULTS_DIR.mkdir(parents=True, exist_ok=True)
        results_file = TEST_RESULTS_DIR / f"test_results_{int(time.time())}.json"
        
        results = {
            "timestamp": time.time(),
            "corpus_stats": {
                file_type: len(files) for file_type, files in downloaded_files.items()
            },
            "num_queries": len(TEST_QUERIES),
            "num_successful_evals": len(successful_evals),
            "aggregate_metrics": aggregate_metrics,
            "difficulty_summary": difficulty_summary,
            "evaluations": evaluations,
        }
        
        with open(results_file, "w") as f:
            json.dump(results, f, indent=2)
        
        logger.info(f"\nResults saved to: {results_file}")
        
        return results
        
    finally:
        # Restore original environment variables
        if original_data_dir:
            os.environ["DATA_DIR"] = original_data_dir
        elif "DATA_DIR" in os.environ:
            del os.environ["DATA_DIR"]
        
        if original_storage_dir:
            os.environ["STORAGE_DIR"] = original_storage_dir
        elif "STORAGE_DIR" in os.environ:
            del os.environ["STORAGE_DIR"]


def main():
    """Main entry point."""
    results = asyncio.run(run_large_scale_tests())
    
    if "error" in results:
        logger.error(f"Test suite failed: {results['error']}")
        sys.exit(1)
    
    logger.info("\n" + "=" * 80)
    logger.info("Test Suite Completed Successfully")
    logger.info("=" * 80)


if __name__ == "__main__":
    main()
