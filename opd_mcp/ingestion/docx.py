"""DOCX document processing."""

import logging
from datetime import datetime
from pathlib import Path
from typing import List

from docx import Document
from llama_index.core import Document as LlamaIndexDocument

from opd_mcp.config import EXCLUDED_EMBED_METADATA_KEYS, EXCLUDED_LLM_METADATA_KEYS
from opd_mcp.storage import get_heading_store

logger = logging.getLogger(__name__)


def _parse_heading_level(style_name: str | None) -> int:
    """Best-effort extraction of a numeric heading level from a DOCX style name."""
    if not style_name:
        return 1
    try:
        if "Heading" in style_name:
            level_str = style_name.replace("Heading", "").strip()
            if level_str:
                return int(level_str)
    except (ValueError, AttributeError):
        pass
    return 1


def split_docx_into_heading_documents(docx_path: Path) -> List[LlamaIndexDocument]:
    """Split DOCX into documents by heading.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        List of LlamaIndexDocument objects, one per heading section
    """
    docs: List[LlamaIndexDocument] = []
    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.warning(f"Failed to open DOCX {docx_path}: {e}")
        return docs

    # Extract file dates from filesystem
    stat = docx_path.stat()
    creation_date = datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d")
    last_modified_date = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d")

    # Try to extract dates from DOCX core properties (more accurate than filesystem)
    try:
        core_props = doc.core_properties
        if core_props.created:
            creation_date = core_props.created.strftime("%Y-%m-%d")
        if core_props.modified:
            last_modified_date = core_props.modified.strftime("%Y-%m-%d")
    except Exception:
        pass  # Fall back to filesystem dates

    # First pass: Extract all headings with positions for hierarchy metadata
    all_headings = []
    char_position = 0
    for para in doc.paragraphs:
        style_name = getattr(para.style, "name", "") or ""
        is_heading = (
            style_name.startswith("Heading")
            or style_name.startswith("heading")
            or "Heading" in style_name
        )

        if is_heading and para.text.strip():
            heading_level = _parse_heading_level(style_name)
            all_headings.append(
                {"text": para.text.strip(), "position": char_position, "level": heading_level}
            )

        char_position += len(para.text) + 1  # +1 for newline

    # Store headings separately to avoid metadata size issues during chunking
    get_heading_store().set_headings(str(docx_path), all_headings)

    # Second pass: Split by heading
    current_heading: str | None = None
    current_level: int | None = None
    current_body: list[str] = []

    def flush_current():
        if not current_heading:
            return
        text = "\n".join(line for line in current_body if line is not None).strip()
        if not text:
            return

        # Build hierarchical heading_path by finding parent headings based on level
        heading_path = []
        if all_headings:
            # Find the index of the current heading in all_headings
            current_idx = None
            for idx, h in enumerate(all_headings):
                if h["text"] == current_heading and h["level"] == current_level:
                    current_idx = idx
                    break

            if current_idx is not None:
                # Build path by including all parent headings
                path_headings = [all_headings[current_idx]]
                for idx in range(current_idx - 1, -1, -1):
                    h = all_headings[idx]
                    if h["level"] < path_headings[0]["level"]:
                        path_headings.insert(0, h)
                heading_path = [h["text"] for h in path_headings]

        metadata = {
            "file_path": str(docx_path),
            "file_name": docx_path.name,
            "source": str(docx_path),
            "heading": current_heading,
            "heading_level": current_level,
            "creation_date": creation_date,
            "last_modified_date": last_modified_date,
            "heading_path": heading_path,
        }
        docs.append(
            LlamaIndexDocument(
                text=text,
                metadata=metadata,
                excluded_embed_metadata_keys=EXCLUDED_EMBED_METADATA_KEYS,
                excluded_llm_metadata_keys=EXCLUDED_LLM_METADATA_KEYS,
            )
        )

    for para in doc.paragraphs:
        style_name = getattr(para.style, "name", "") or ""
        is_heading = (
            style_name.startswith("Heading")
            or style_name.startswith("heading")
            or "Heading" in style_name
        )

        if is_heading and para.text.strip():
            flush_current()
            current_heading = para.text.strip()
            current_level = _parse_heading_level(style_name)
            current_body = []
        else:
            if current_heading is not None:
                current_body.append(para.text)

    flush_current()

    if not docs:
        try:
            full_text = "\n".join(p.text for p in doc.paragraphs).strip()
        except Exception:
            full_text = ""

        if full_text:
            metadata = {
                "file_path": str(docx_path),
                "file_name": docx_path.name,
                "source": str(docx_path),
                "heading": None,
                "heading_level": None,
                "creation_date": creation_date,
                "last_modified_date": last_modified_date,
            }
            docs.append(
                LlamaIndexDocument(
                    text=full_text,
                    metadata=metadata,
                    excluded_embed_metadata_keys=EXCLUDED_EMBED_METADATA_KEYS,
                    excluded_llm_metadata_keys=EXCLUDED_LLM_METADATA_KEYS,
                )
            )

    logger.info(f"Split DOCX {docx_path} into {len(docs)} heading-based document(s)")
    return docs
