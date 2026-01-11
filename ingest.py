#!/usr/bin/env python3
"""
Ingestion pipeline for building a RAG index from PDF, DOCX, Markdown, and TXT documents.
Supports incremental indexing using a local SQLite database to track file states.
"""
import argparse
import hashlib
import itertools
import json
import logging
import os
import sqlite3
import sys
import threading
import time
from datetime import datetime
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Optional, Iterator, Set

# Check for --offline flag early and set environment variables BEFORE any imports
if "--offline" in sys.argv:
    os.environ["HF_HUB_OFFLINE"] = "1"
    os.environ["TRANSFORMERS_OFFLINE"] = "1"
    os.environ["HF_DATASETS_OFFLINE"] = "1"

from docx import Document

from llama_index.core import (
    VectorStoreIndex,
    StorageContext,
    Settings,
    SimpleDirectoryReader,
    Document as LlamaIndexDocument,
    load_index_from_storage,
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.fastembed import FastEmbedEmbedding

# Load environment variables

# Configuration
DATA_DIR = Path(os.getenv("DATA_DIR", "./data"))
STORAGE_DIR = Path(os.getenv("STORAGE_DIR", "./storage"))
STATE_DB_PATH = STORAGE_DIR / "ingestion_state.db"

# Stage 1 (embedding/vector search) configuration
RETRIEVAL_EMBED_MODEL_NAME = os.getenv(
    "RETRIEVAL_EMBED_MODEL_NAME", "BAAI/bge-small-en-v1.5"
)

# Stage 2 (FlashRank reranking, CPU-only, ONNX-based) configuration
RETRIEVAL_RERANK_MODEL_NAME = os.getenv(
    "RETRIEVAL_RERANK_MODEL_NAME", "ms-marco-MiniLM-L-12-v2"
)

# Shared cache directory for embedding and reranking models
RETRIEVAL_MODEL_CACHE_DIR = Path(os.getenv("RETRIEVAL_MODEL_CACHE_DIR", "./models"))

# Text chunking configuration
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "512"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "100"))

# BM25 index directory for file name matching
BM25_INDEX_DIR = STORAGE_DIR / "bm25_index"

# Heading store for document headings (stored separately to avoid metadata size issues)
HEADING_STORE_PATH = STORAGE_DIR / "heading_store.json"

# Metadata exclusion configuration
# These keys are excluded from the embedding text to save tokens and avoid length errors
EXCLUDED_EMBED_METADATA_KEYS = [
    "line_offsets",      # Large integer array, primary cause of length errors
    "document_headings", # Heading hierarchy array with positions, excluded like line_offsets
    "heading_path",      # Pre-computed heading hierarchy, stored separately to save chunk space
    "file_path",         # redundant with file_name/source, strict path less useful for semantic similarity
    "source",            # often same as file_path
    "creation_date",     # temporal, not semantic
    "last_modified_date",# temporal, not semantic
    "doc_ids",           # internal tracking
    "hash",              # internal tracking
]

# These keys are excluded from the LLM context to save context window
EXCLUDED_LLM_METADATA_KEYS = [
    "line_offsets",      # LLM needs text content, not integer map
    "hash",              # internal tracking
    "doc_ids",           # internal tracking
    "file_path",         # usually redundant if file_name is present
    "source",            # usually redundant
]

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


class HeadingStore:
    """Stores document headings separately from chunk metadata.

    This avoids the LlamaIndex SentenceSplitter metadata size validation issue,
    which checks metadata length before applying exclusions. By storing headings
    in a separate file, we keep chunk metadata small while preserving heading
    data for retrieval.
    """

    def __init__(self, store_path: Path):
        self.store_path = store_path
        self._data: Dict[str, List[dict]] = {}
        self._load()

    def _load(self):
        """Load heading data from disk."""
        if self.store_path.exists():
            try:
                with open(self.store_path, "r", encoding="utf-8") as f:
                    self._data = json.load(f)
            except Exception as e:
                logger.warning(f"Failed to load heading store: {e}")
                self._data = {}

    def _save(self):
        """Save heading data to disk."""
        self.store_path.parent.mkdir(parents=True, exist_ok=True)
        with open(self.store_path, "w", encoding="utf-8") as f:
            json.dump(self._data, f)

    def set_headings(self, file_path: str, headings: List[dict]):
        """Store headings for a file."""
        self._data[file_path] = headings
        self._save()

    def get_headings(self, file_path: str) -> List[dict]:
        """Get headings for a file."""
        return self._data.get(file_path, [])

    def remove_headings(self, file_path: str):
        """Remove headings for a file."""
        if file_path in self._data:
            del self._data[file_path]
            self._save()


# Module-level heading store instance (lazy initialized)
_heading_store: Optional["HeadingStore"] = None


def get_heading_store() -> HeadingStore:
    """Get the singleton HeadingStore instance."""
    global _heading_store
    if _heading_store is None:
        _heading_store = HeadingStore(HEADING_STORE_PATH)
    return _heading_store


@dataclass
class FileInfo:
    """Metadata about a file in the data source."""
    path: str
    hash: str
    last_modified: float


class IngestionState:
    """Manages the state of ingested files using a SQLite database."""

    def __init__(self, db_path: Path):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        """Initialize the SQLite database schema."""
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        with sqlite3.connect(self.db_path) as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS files (
                    path TEXT PRIMARY KEY,
                    hash TEXT NOT NULL,
                    last_modified REAL NOT NULL,
                    doc_ids TEXT NOT NULL
                )
                """
            )

    def get_all_files(self) -> Dict[str, dict]:
        """Retrieve all tracked files and their metadata."""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute("SELECT path, hash, last_modified, doc_ids FROM files")
            return {
                row[0]: {
                    "hash": row[1],
                    "last_modified": row[2],
                    "doc_ids": row[3].split(",") if row[3] else [],
                }
                for row in cursor
            }

    def update_file_state(self, file_info: FileInfo, doc_ids: List[str]):
        """Update or insert the state for a file."""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute(
                """
                INSERT INTO files (path, hash, last_modified, doc_ids)
                VALUES (?, ?, ?, ?)
                ON CONFLICT(path) DO UPDATE SET
                    hash=excluded.hash,
                    last_modified=excluded.last_modified,
                    doc_ids=excluded.doc_ids
                """,
                (
                    file_info.path,
                    file_info.hash,
                    file_info.last_modified,
                    ",".join(doc_ids),
                ),
            )

    def remove_file_state(self, path: str):
        """Remove a file from the state tracking."""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("DELETE FROM files WHERE path = ?", (path,))


class DataSource(ABC):
    """Abstract base class for data sources."""

    @abstractmethod
    def iter_files(self) -> Iterator[FileInfo]:
        """Yield FileInfo for each file in the source."""
        pass

    @abstractmethod
    def load_file(self, file_info: FileInfo) -> List[LlamaIndexDocument]:
        """Load and return documents for a given file."""
        pass


def _compute_line_offsets(text: str) -> List[int]:
    """Compute character offset positions for each line start.

    Returns a list where line_offsets[i] is the character position where line i+1 starts.
    Line 1 starts at position 0 (implicit).
    """
    offsets = [0]  # Line 1 starts at position 0
    for i, char in enumerate(text):
        if char == '\n':
            offsets.append(i + 1)  # Next line starts after the newline
    return offsets


def _extract_markdown_headings(text: str) -> List[dict]:
    """Extract heading hierarchy from Markdown text using ATX-style syntax.

    Parses # Heading syntax and returns list of dicts with text, position, level.
    Handles ATX-style headings (# Heading) but not Setext (underlined).

    Returns:
        List of dicts with keys: text (str), position (int), level (int)
    """
    import re

    headings = []
    # Match ATX-style headings: line start, 1-6 #s, space, text
    pattern = re.compile(r'^(#{1,6})\s+(.+?)$', re.MULTILINE)

    # Find all code block ranges to skip headings inside them
    code_blocks = []
    for match in re.finditer(r'```.*?```', text, flags=re.DOTALL):
        code_blocks.append((match.start(), match.end()))

    def is_in_code_block(pos):
        """Check if position is inside a code block."""
        return any(start <= pos < end for start, end in code_blocks)

    for match in pattern.finditer(text):
        # Skip headings inside code blocks
        if is_in_code_block(match.start()):
            continue

        level = len(match.group(1))
        heading_text = match.group(2).strip()
        position = match.start()

        if heading_text:
            headings.append({
                "text": heading_text,
                "position": position,
                "level": level
            })

    return headings


def _extract_pdf_headings_from_outline(pdf_path: Path) -> List[dict]:
    """Extract headings from PDF outline/bookmarks (TOC).

    Returns list of dicts with text, position (estimated), level.
    Position is approximate based on cumulative page character counts.
    Falls back to empty list if PDF has no outline or extraction fails.

    Returns:
        List of dicts with keys: text (str), position (int), level (int)
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        logger.warning("PyMuPDF not available, skipping PDF heading extraction")
        return []

    try:
        doc = fitz.open(pdf_path)
        toc = doc.get_toc()  # Returns [[level, title, page_num], ...]

        if not toc:
            return []

        headings = []
        for item in toc:
            level, title, page_num = item[0], item[1], item[2]

            # Estimate position by accumulating text from previous pages
            position = 0
            for page_idx in range(page_num - 1):
                if page_idx < len(doc):
                    page = doc[page_idx]
                    position += len(page.get_text())

            headings.append({
                "text": title.strip(),
                "position": position,
                "level": level
            })

        doc.close()
        return headings

    except Exception as e:
        logger.warning(f"Failed to extract PDF outline from {pdf_path}: {e}")
        return []


class LocalFileSystemSource(DataSource):
    """Implementation of DataSource for the local file system."""

    def __init__(self, base_dir: Path):
        self.base_dir = base_dir

    def iter_files(self) -> Iterator[FileInfo]:
        extensions = {".pdf", ".md", ".txt", ".docx"}
        for root, _, files in os.walk(self.base_dir):
            for file in files:
                if Path(file).suffix.lower() in extensions:
                    file_path = Path(root) / file
                    yield self._create_file_info(file_path)

    def _create_file_info(self, file_path: Path) -> FileInfo:
        # Create a hash of the file content
        hasher = hashlib.md5()
        with open(file_path, "rb") as f:
            buf = f.read(65536)
            while len(buf) > 0:
                hasher.update(buf)
                buf = f.read(65536)
        
        return FileInfo(
            path=str(file_path.absolute()),
            hash=hasher.hexdigest(),
            last_modified=file_path.stat().st_mtime,
        )

    def load_file(self, file_info: FileInfo) -> List[LlamaIndexDocument]:
        file_path = Path(file_info.path)
        if file_path.suffix.lower() == ".docx":
            return split_docx_into_heading_documents(file_path)
        else:
            reader = SimpleDirectoryReader(
                input_files=[str(file_path)],
            )
            docs = reader.load_data()
            # Ensure dates are visible to LLM (remove from exclusion list)
            for doc in docs:
                if hasattr(doc, 'excluded_llm_metadata_keys') and doc.excluded_llm_metadata_keys:
                    doc.excluded_llm_metadata_keys = [
                        k for k in doc.excluded_llm_metadata_keys
                        if k not in ('creation_date', 'last_modified_date')
                    ]

            # Add line offsets for text-based files (markdown, txt) to enable line number lookup
            if file_path.suffix.lower() in {".md", ".txt"}:
                for doc in docs:
                    text = doc.get_content()
                    line_offsets = _compute_line_offsets(text)
                    doc.metadata["line_offsets"] = line_offsets

                    # Extract headings for Markdown and store separately
                    # (not in metadata to avoid SentenceSplitter size validation)
                    if file_path.suffix.lower() == ".md":
                        headings = _extract_markdown_headings(text)
                        get_heading_store().set_headings(str(file_path), headings)

            # Extract headings for PDF files and store separately
            if file_path.suffix.lower() == ".pdf":
                headings = _extract_pdf_headings_from_outline(file_path)
                get_heading_store().set_headings(str(file_path), headings)

            # Apply metadata exclusions
            for doc in docs:
                doc.excluded_embed_metadata_keys = EXCLUDED_EMBED_METADATA_KEYS
                doc.excluded_llm_metadata_keys = EXCLUDED_LLM_METADATA_KEYS

            return docs


class SimpleProgressBar:
    """Lightweight progress bar using only the standard library."""

    def __init__(self, total: int, desc: str, unit: str = "item", width: int = 30):
        self.total = max(total, 0)
        self.desc = desc
        self.unit = unit
        self.width = width
        self.current = 0
        if self.total > 0:
            self._render()

    def update(self, step: int = 1) -> None:
        if self.total <= 0:
            return
        self.current = min(self.total, self.current + step)
        self._render()
        if self.current >= self.total:
            sys.stdout.write("\n")
            sys.stdout.flush()

    def _render(self) -> None:
        progress = self.current / self.total if self.total else 0
        filled = int(self.width * progress)
        bar = "#" * filled + "-" * (self.width - filled)
        sys.stdout.write(
            f"\r{self.desc} [{bar}] {progress * 100:5.1f}% ({self.current}/{self.total} {self.unit}s)"
        )
        sys.stdout.flush()


class Spinner:
    """Simple console spinner to indicate long-running steps."""

    def __init__(self, desc: str, interval: float = 0.1):
        self.desc = desc
        self.interval = interval
        self._stop_event = threading.Event()
        self._thread: threading.Thread | None = None
        self._line = desc

    def __enter__(self):
        self._thread = threading.Thread(target=self._spin, daemon=True)
        self._thread.start()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self._stop_event.set()
        if self._thread:
            self._thread.join()
        # Clear spinner line
        sys.stdout.write("\r" + " " * len(self._line) + "\r")
        sys.stdout.flush()

    def _spin(self) -> None:
        for char in itertools.cycle("|/-\\"):
            if self._stop_event.is_set():
                break
            self._line = f"{self.desc} {char}"
            sys.stdout.write("\r" + self._line)
            sys.stdout.flush()
            time.sleep(self.interval)


def _embedding_cache_path(model_name: str, cache_dir: Path) -> Path:
    """Return the expected cache directory for a FastEmbed model."""
    return cache_dir / f"models--{model_name.replace('/', '--')}"


def _verify_model_cache_exists(cache_dir: Path) -> bool:
    """
    Verify that the cached model directory exists and contains the expected model files.
    """
    from fastembed import TextEmbedding
    
    try:
        models = TextEmbedding.list_supported_models()
        model_info = [m for m in models if m.get("model") == RETRIEVAL_EMBED_MODEL_NAME]
        if not model_info:
            return False
        
        model_info = model_info[0]
        hf_source = model_info.get("sources", {}).get("hf")
        if not hf_source:
            return False
        
        expected_dir = cache_dir / f"models--{hf_source.replace('/', '--')}"
        if not expected_dir.exists():
            return False
        
        snapshots_dir = expected_dir / "snapshots"
        if not snapshots_dir.exists():
            return False
        
        model_file = model_info.get("model_file", "model_optimized.onnx")
        for snapshot in snapshots_dir.iterdir():
            if snapshot.is_dir():
                model_path = snapshot / model_file
                if model_path.exists() or model_path.is_symlink():
                    return True
        
        return False
    except Exception:
        return False


def _get_cached_model_path(cache_dir: Path, model_name: str) -> Path | None:
    """Get the cached model directory path."""
    try:
        from huggingface_hub import snapshot_download
        from fastembed import TextEmbedding
        models = TextEmbedding.list_supported_models()
        model_info = [m for m in models if m.get("model") == model_name]
        if model_info:
            hf_source = model_info[0].get("sources", {}).get("hf")
            if hf_source:
                cache_dir_abs = cache_dir.resolve()
                model_dir = snapshot_download(
                    repo_id=hf_source,
                    local_files_only=True,
                    cache_dir=str(cache_dir_abs)
                )
                return Path(model_dir).resolve()
    except (ImportError, Exception):
        pass
    return None


def _create_fastembed_embedding(cache_dir: Path, offline: bool = False):
    """Create a FastEmbedEmbedding instance."""
    if offline:
        cached_model_path = _get_cached_model_path(cache_dir, RETRIEVAL_EMBED_MODEL_NAME)
        if cached_model_path:
            logger.info(
                f"Using cached model path to bypass download: {cached_model_path}"
            )
            return FastEmbedEmbedding(
                model_name=RETRIEVAL_EMBED_MODEL_NAME,
                cache_dir=str(cache_dir),
                specific_model_path=str(cached_model_path)
            )
        else:
            logger.warning(
                "Could not find cached model path, falling back to normal initialization"
            )
    
    return FastEmbedEmbedding(
        model_name=RETRIEVAL_EMBED_MODEL_NAME, cache_dir=str(cache_dir)
    )


def ensure_embedding_model_cached(cache_dir: Path, offline: bool = False) -> None:
    """Ensure the embedding model is available in the local cache."""
    if offline:
        logger.info("Verifying embedding model cache...")
        if _verify_model_cache_exists(cache_dir):
            logger.info("Embedding model found in cache")
        else:
            logger.error(
                "Offline mode enabled, but embedding model cache not found in %s",
                cache_dir,
            )
            raise FileNotFoundError(
                f"Embedding model '{RETRIEVAL_EMBED_MODEL_NAME}' not found in cache directory '{cache_dir}'. "
            )
    
    try:
        logger.info("Initializing embedding model from cache...")
        cache_dir_abs = cache_dir.resolve()
        if offline:
            os.environ["HF_HUB_CACHE"] = str(cache_dir_abs)
        
        _create_fastembed_embedding(cache_dir, offline=offline)
        logger.info("Embedding model initialized successfully")
        return
    except (ValueError, Exception) as e:
        # Simplified error handling for brevity, similar logic as original
        if offline:
            raise FileNotFoundError(f"Failed to load model offline: {e}") from e
        else:
            raise RuntimeError(f"Failed to download/initialize model: {e}") from e


def ensure_rerank_model_cached(cache_dir: Path, offline: bool = False) -> Path:
    """Ensure the reranking model is cached locally."""
    try:
        from flashrank import Ranker
    except ImportError as exc:
        raise ImportError(
            "flashrank is required for reranking."
        ) from exc

    cache_dir_abs = cache_dir.resolve()
    logger.info("Ensuring rerank model is available in cache...")

    # Map cross-encoder model names to FlashRank equivalents if needed
    model_name = RETRIEVAL_RERANK_MODEL_NAME
    # Note: FlashRank doesn't have L-6 models, so we map to L-12 equivalents
    model_mapping = {
        "cross-encoder/ms-marco-MiniLM-L-6-v2": "ms-marco-MiniLM-L-12-v2",  # L-6 not available, use L-12
        "ms-marco-MiniLM-L-6-v2": "ms-marco-MiniLM-L-12-v2",  # Direct mapping for L-6
    }
    if model_name in model_mapping:
        model_name = model_mapping[model_name]
    elif model_name.startswith("cross-encoder/"):
        # Extract model name after cross-encoder/ prefix and try to map
        base_name = model_name.replace("cross-encoder/", "")
        # If it's an L-6 model, map to L-12
        if "L-6" in base_name:
            model_name = base_name.replace("L-6", "L-12")
        else:
            model_name = base_name

    try:
        reranker = Ranker(model_name=model_name, cache_dir=str(cache_dir_abs))
        logger.info(f"FlashRank model '{model_name}' initialized successfully")
        return cache_dir_abs
    except Exception as exc:
        if offline:
            raise FileNotFoundError(
                f"Rerank model '{model_name}' not found in cache."
            ) from exc
        raise


def _parse_heading_level(style_name: str | None) -> int:
    """Best-effort extraction of a numeric heading level from a DOCX style name."""
    if not style_name:
        return 1
    try:
        if "Heading" in style_name:
            level_str = style_name.replace("Heading", "").strip()
            if level_str:
                return int(level_str)
    except (ValueError, AttributeError):
        pass
    return 1


def split_docx_into_heading_documents(docx_path: Path) -> List[LlamaIndexDocument]:
    """Split DOCX into documents by heading."""
    docs: List[LlamaIndexDocument] = []
    try:
        doc = Document(docx_path)
    except Exception as e:
        logger.warning(f"Failed to open DOCX {docx_path}: {e}")
        return docs

    # Extract file dates from filesystem
    stat = docx_path.stat()
    creation_date = datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d")
    last_modified_date = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d")

    # Try to extract dates from DOCX core properties (more accurate than filesystem)
    try:
        core_props = doc.core_properties
        if core_props.created:
            creation_date = core_props.created.strftime("%Y-%m-%d")
        if core_props.modified:
            last_modified_date = core_props.modified.strftime("%Y-%m-%d")
    except Exception:
        pass  # Fall back to filesystem dates

    # First pass: Extract all headings with positions for hierarchy metadata
    all_headings = []
    char_position = 0
    for para in doc.paragraphs:
        style_name = getattr(para.style, "name", "") or ""
        is_heading = (
            style_name.startswith("Heading")
            or style_name.startswith("heading")
            or "Heading" in style_name
        )

        if is_heading and para.text.strip():
            heading_level = _parse_heading_level(style_name)
            all_headings.append({
                "text": para.text.strip(),
                "position": char_position,
                "level": heading_level
            })

        char_position += len(para.text) + 1  # +1 for newline

    # Store headings separately to avoid metadata size issues during chunking
    get_heading_store().set_headings(str(docx_path), all_headings)

    # Second pass: Split by heading (existing logic)
    current_heading: str | None = None
    current_level: int | None = None
    current_body: list[str] = []

    def flush_current():
        if not current_heading:
            return
        text = "\n".join(line for line in current_body if line is not None).strip()
        if not text:
            return

        # Build hierarchical heading_path by finding parent headings based on level
        heading_path = []
        if all_headings:
            # Find the index of the current heading in all_headings
            current_idx = None
            for idx, h in enumerate(all_headings):
                if h["text"] == current_heading and h["level"] == current_level:
                    current_idx = idx
                    break

            if current_idx is not None:
                # Build path by including all parent headings (those with lower level numbers)
                # Walk backwards from current heading and include headings with level < current_level
                path_headings = [all_headings[current_idx]]  # Start with current
                for idx in range(current_idx - 1, -1, -1):
                    h = all_headings[idx]
                    if h["level"] < path_headings[0]["level"]:
                        path_headings.insert(0, h)
                heading_path = [h["text"] for h in path_headings]

        metadata = {
            "file_path": str(docx_path),
            "file_name": docx_path.name,
            "source": str(docx_path),
            "heading": current_heading,
            "heading_level": current_level,
            "creation_date": creation_date,
            "last_modified_date": last_modified_date,
            "heading_path": heading_path,  # Pre-computed hierarchical path
        }
        docs.append(LlamaIndexDocument(
            text=text,
            metadata=metadata,
            excluded_embed_metadata_keys=EXCLUDED_EMBED_METADATA_KEYS,
            excluded_llm_metadata_keys=EXCLUDED_LLM_METADATA_KEYS,
        ))

    for para in doc.paragraphs:
        style_name = getattr(para.style, "name", "") or ""
        is_heading = (
            style_name.startswith("Heading")
            or style_name.startswith("heading")
            or "Heading" in style_name
        )

        if is_heading and para.text.strip():
            flush_current()
            current_heading = para.text.strip()
            current_level = _parse_heading_level(style_name)
            current_body = []
        else:
            if current_heading is not None:
                current_body.append(para.text)

    flush_current()

    if not docs:
        try:
            full_text = "\n".join(p.text for p in doc.paragraphs).strip()
        except Exception:
            full_text = ""

        if full_text:
            metadata = {
                "file_path": str(docx_path),
                "file_name": docx_path.name,
                "source": str(docx_path),
                "heading": None,
                "heading_level": None,
                "creation_date": creation_date,
                "last_modified_date": last_modified_date,
            }
            docs.append(LlamaIndexDocument(
                text=full_text,
                metadata=metadata,
                excluded_embed_metadata_keys=EXCLUDED_EMBED_METADATA_KEYS,
                excluded_llm_metadata_keys=EXCLUDED_LLM_METADATA_KEYS,
            ))

    logger.info(
        f"Split DOCX {docx_path} into {len(docs)} heading-based document(s)"
    )
    return docs


def tokenize_filename(filename: str) -> List[str]:
    """
    Tokenize a filename for BM25 indexing.

    Splits on delimiters (underscore, hyphen, dot, space) and camelCase.

    Examples:
        'cpp_styleguide.md' -> ['cpp', 'styleguide', 'md']
        'API-Reference-v2.pdf' -> ['api', 'reference', 'v2', 'pdf']
        'CamelCaseDoc.docx' -> ['camel', 'case', 'doc', 'docx']
    """
    import re

    name_parts = filename.rsplit('.', 1)
    base_name = name_parts[0]
    extension = name_parts[1] if len(name_parts) > 1 else ""

    # Split on explicit delimiters
    parts = re.split(r'[_\-\.\s]+', base_name)

    # Split camelCase within each part
    tokens = []
    for part in parts:
        camel_split = re.sub(r'([a-z])([A-Z])', r'\1 \2', part).split()
        tokens.extend(t.lower() for t in camel_split if t)

    # Add extension as a token
    if extension:
        tokens.append(extension.lower())

    return tokens


def build_bm25_index(index, storage_dir: Path) -> None:
    """
    Build a BM25 index over file names from the docstore.

    This enables keyword matching for queries like 'cpp styleguide' to find
    files named 'cpp_styleguide.md'.
    """
    from llama_index.retrievers.bm25 import BM25Retriever
    from llama_index.core.schema import TextNode

    logger.info("Building BM25 index for file name matching...")

    # Create filename nodes - one per unique file
    filename_nodes = []
    seen_files: Set[str] = set()

    for doc_id, node in index.docstore.docs.items():
        metadata = node.metadata or {}
        file_name = metadata.get("file_name", "")
        file_path = metadata.get("file_path", "")

        if not file_name or file_path in seen_files:
            continue
        seen_files.add(file_path)

        tokens = tokenize_filename(file_name)
        filename_nodes.append(TextNode(
            text=" ".join(tokens),
            metadata={"file_name": file_name, "file_path": file_path},
            id_=f"bm25_{file_path}"
        ))

    if not filename_nodes:
        logger.warning("No documents found for BM25 indexing")
        return

    logger.info(f"Creating BM25 index with {len(filename_nodes)} file name entries")

    bm25_retriever = BM25Retriever.from_defaults(
        nodes=filename_nodes,
        similarity_top_k=10,
    )

    bm25_dir = storage_dir / "bm25_index"
    bm25_dir.mkdir(parents=True, exist_ok=True)
    bm25_retriever.persist(str(bm25_dir))

    logger.info(f"BM25 index persisted to {bm25_dir}")


def configure_offline_mode(offline: bool, cache_dir: Path) -> None:
    """Configure environment variables for offline mode."""
    if offline:
        os.environ["HF_HUB_OFFLINE"] = "1"
        os.environ["TRANSFORMERS_OFFLINE"] = "1"
        os.environ["HF_DATASETS_OFFLINE"] = "1"
        cache_dir_abs = cache_dir.resolve()
        os.environ["HF_HOME"] = str(cache_dir_abs)
        os.environ["HF_HUB_CACHE"] = str(cache_dir_abs)
        os.environ["HF_DATASETS_CACHE"] = str(cache_dir_abs)
        logger.info("Offline mode enabled.")
    else:
        # Clear offline mode environment variables to allow downloads
        for var in ["HF_HUB_OFFLINE", "TRANSFORMERS_OFFLINE", "HF_DATASETS_OFFLINE"]:
            os.environ.pop(var, None)

    # Update huggingface_hub's cached constant (it caches at import time)
    try:
        from huggingface_hub import constants
        constants.HF_HUB_OFFLINE = offline
    except ImportError:
        pass


def build_index(download_only: bool = False, offline: bool = False) -> None:
    """Build and persist the vector index incrementally."""
    configure_offline_mode(offline, RETRIEVAL_MODEL_CACHE_DIR)
    
    logger.info(f"Starting ingestion from {DATA_DIR}")

    ensure_embedding_model_cached(RETRIEVAL_MODEL_CACHE_DIR, offline=offline)
    try:
        ensure_rerank_model_cached(RETRIEVAL_MODEL_CACHE_DIR, offline=offline)
    except FileNotFoundError:
        if download_only or offline:
            raise
        logger.warning("Rerank model could not be cached yet; continuing without it.")

    if download_only:
        logger.info("Models downloaded; skipping index build.")
        return

    # Initialize State and Data Source
    ingestion_state = IngestionState(STATE_DB_PATH)
    data_source = LocalFileSystemSource(DATA_DIR)

    # Initialize Embedding Model
    logger.info(f"Initializing embedding model: {RETRIEVAL_EMBED_MODEL_NAME}")
    with Spinner("Initializing embedding model"):
        embed_model = _create_fastembed_embedding(RETRIEVAL_MODEL_CACHE_DIR, offline=offline)
    Settings.embed_model = embed_model
    
    # Configure Text Splitter
    text_splitter = SentenceSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separator=" ",
    )
    Settings.text_splitter = text_splitter

    # Load existing index or create new
    if (STORAGE_DIR / "docstore.json").exists():
        logger.info("Loading existing index context...")
        storage_context = StorageContext.from_defaults(persist_dir=str(STORAGE_DIR))
        index = load_index_from_storage(storage_context, embed_model=embed_model)
    else:
        logger.info("Creating new index context...")
        storage_context = StorageContext.from_defaults()
        index = VectorStoreIndex([], storage_context=storage_context, embed_model=embed_model)

    # Change Detection
    tracked_files = ingestion_state.get_all_files()
    found_files: Set[str] = set()
    files_to_process: List[FileInfo] = []

    logger.info("Scanning for changes...")
    for file_info in data_source.iter_files():
        found_files.add(file_info.path)
        existing_state = tracked_files.get(file_info.path)
        
        if existing_state:
            # Check if modified
            if existing_state["hash"] != file_info.hash:
                logger.info(f"Modified file detected: {file_info.path}")
                files_to_process.append(file_info)
        else:
            # New file
            logger.info(f"New file detected: {file_info.path}")
            files_to_process.append(file_info)

    # Identify Deleted Files
    deleted_files = set(tracked_files.keys()) - found_files
    for deleted_path in deleted_files:
        logger.info(f"Deleted file detected: {deleted_path}")
        doc_ids = tracked_files[deleted_path]["doc_ids"]
        for doc_id in doc_ids:
            try:
                index.delete_ref_doc(doc_id, delete_from_docstore=True)
            except Exception as e:
                logger.warning(f"Failed to delete doc {doc_id} from index: {e}")
        # Clean up heading data for deleted file
        get_heading_store().remove_headings(deleted_path)
        ingestion_state.remove_file_state(deleted_path)

    if not files_to_process and not deleted_files:
        logger.info("No changes detected. Index is up to date.")
        return

    # Process New/Modified Files
    if files_to_process:
        progress = SimpleProgressBar(len(files_to_process), desc="Processing files", unit="file")
        for file_info in files_to_process:
            # Remove old versions if they exist
            existing_state = tracked_files.get(file_info.path)
            if existing_state:
                for doc_id in existing_state["doc_ids"]:
                    try:
                        index.delete_ref_doc(doc_id, delete_from_docstore=True)
                    except KeyError:
                        pass # Document might already be gone

            # Load and Index New Version
            docs = data_source.load_file(file_info)
            doc_ids = []
            for doc in docs:
                index.insert(doc)
                doc_ids.append(doc.doc_id)
            
            # Update State
            ingestion_state.update_file_state(file_info, doc_ids)
            progress.update()

    # Persist Index
    STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    logger.info(f"Persisting index to {STORAGE_DIR}")
    index.storage_context.persist(persist_dir=str(STORAGE_DIR))

    # Build BM25 index for file name matching
    build_bm25_index(index, STORAGE_DIR)

    logger.info("Ingestion complete.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build the document index")
    parser.add_argument(
        "--download-models",
        action="store_true",
        help="Download the retrieval models and exit",
    )
    parser.add_argument(
        "--offline",
        action="store_true",
        help="Run entirely offline",
    )
    args = parser.parse_args()

    try:
        build_index(download_only=args.download_models, offline=args.offline)
    except Exception as e:
        logger.error(f"Ingestion failed: {e}", exc_info=True)
        raise


